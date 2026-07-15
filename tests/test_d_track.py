from __future__ import annotations

import json
from pathlib import Path
from types import SimpleNamespace

import pytest
from PIL import Image
from docx import Document
from pypdf import PdfWriter

from app.delivery_validator import DeliveryValidator
from app.exporter_base import ExportBaseMixin
from app.exporter_render import ExportRenderMixin
from app.figure_protocol import (
    FigureProtocolError,
    artifact_reference,
    parse_figure_block,
    resolve_figure_reference,
)
from app.pdf_exporter import PdfConverter, PdfConversionError


def test_consecutive_figure_directives_are_not_merged(tmp_path: Path):
    first = tmp_path / "first.png"
    second = tmp_path / "second.png"
    Image.new("RGB", (800, 400), "white").save(first)
    Image.new("RGB", (800, 400), "white").save(second)
    block = (
        f"[[FIGURE]]artifact://first.png|图1|12|source=artifact://first.mmd\n"
        f"[[FIGURE]]artifact://second.png|图2|13|source=artifact://second.mmd"
    )
    directives = parse_figure_block(block)
    assert [item.caption for item in directives] == ["图1", "图2"]

    renderer = ExportRenderMixin()
    renderer.settings = SimpleNamespace(data_dir=tmp_path)
    document = Document()
    renderer._append_block(document, block)
    assert len(document.inline_shapes) == 2
    assert "first.png" not in "\n".join(p.text for p in document.paragraphs)
    assert "second.png" not in "\n".join(p.text for p in document.paragraphs)


def test_figure_reference_is_portable_and_cannot_escape_data_dir(tmp_path: Path):
    image = tmp_path / "diagram_artifacts" / "p" / "s" / "figure.png"
    image.parent.mkdir(parents=True)
    Image.new("RGB", (400, 300), "white").save(image)
    reference = artifact_reference(image, tmp_path)
    assert reference == "artifact://diagram_artifacts/p/s/figure.png"
    assert resolve_figure_reference(reference, tmp_path) == image.resolve()
    with pytest.raises(FigureProtocolError):
        resolve_figure_reference("artifact://../outside.png", tmp_path)


class _FakeDb:
    def __init__(self, rows):
        self.rows = rows

    def fetchall(self, sql, params):
        return self.rows

    def fetchone(self, sql, params):
        return None


def _run(prompt_id: str, candidate_id: str, created_at: str, *, status: str = "PASS"):
    payload = {
        "source_section": {"section_id": "s1", "title": "第一章", "text_hash": "h"},
        "polished_candidate": {"candidate_id": candidate_id},
    }
    result = {
        "candidate_id": candidate_id,
        "candidate_text": "正文",
        "paragraphs": [{"sequence": 1, "text": "正文"}],
    }
    return {
        "id": f"{prompt_id}-{created_at}",
        "workflow_id": "wf-1",
        "prompt_id": prompt_id,
        "input_json": json.dumps({"payload": payload}, ensure_ascii=False),
        "output_json": json.dumps({"result": result}, ensure_ascii=False),
        "status": status,
        "created_at": created_at,
    }


def test_docx_export_candidates_require_later_expression_critic_pass(tmp_path: Path):
    settings = SimpleNamespace(exports_dir=tmp_path, data_dir=tmp_path)
    approved_rows = [
        _run("P-EXPRESSION-POLISH", "candidate-1", "2026-07-15T00:00:00Z"),
        _run("P-EXPRESSION-CRITIC", "candidate-1", "2026-07-15T00:00:01Z"),
    ]
    exporter = ExportBaseMixin(_FakeDb(approved_rows), settings)
    candidates = exporter._candidate_runs("project-1")
    assert len(candidates) == 1
    assert candidates[0]["candidate_id"] == "candidate-1"
    assert candidates[0]["expression_critic_run_id"].startswith("P-EXPRESSION-CRITIC")

    unapproved = ExportBaseMixin(_FakeDb(approved_rows[:1]), settings)
    assert unapproved._candidate_runs("project-1") == []


def test_pdf_converter_fails_closed_without_libreoffice(tmp_path: Path, monkeypatch):
    docx = tmp_path / "sample.docx"
    Document().save(docx)
    monkeypatch.setattr("app.pdf_exporter.shutil.which", lambda name: None)
    with pytest.raises(PdfConversionError, match="cannot be silently skipped"):
        PdfConverter(SimpleNamespace()).convert(docx)


def test_pdf_converter_writes_hash_log(tmp_path: Path, monkeypatch):
    docx = tmp_path / "sample.docx"
    Document().save(docx)

    def fake_run(command, **kwargs):
        output = docx.with_suffix(".pdf")
        writer = PdfWriter()
        writer.add_blank_page(width=595, height=842)
        with output.open("wb") as handle:
            writer.write(handle)
        return SimpleNamespace(returncode=0, stdout="converted", stderr="")

    monkeypatch.setattr("app.pdf_exporter.shutil.which", lambda name: "/usr/bin/soffice")
    monkeypatch.setattr("app.pdf_exporter.subprocess.run", fake_run)
    pdf = PdfConverter(SimpleNamespace()).convert(docx)
    log = json.loads(docx.with_suffix(".pdf-conversion.json").read_text(encoding="utf-8"))
    assert pdf.exists()
    assert log["status"] == "PASS"
    assert log["docx_sha256"] and log["pdf_sha256"]


def test_structure_validator_detects_empty_section_placeholder_and_raw_marker(tmp_path: Path):
    docx = tmp_path / "bad.docx"
    document = Document()
    document.add_heading("第一章", level=1)
    document.add_heading("第二章", level=1)
    document.add_paragraph("TODO：[[FIGURE]]/tmp/missing.png|图|15")
    document.save(docx)
    pdf = tmp_path / "bad.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=595, height=842)
    with pdf.open("wb") as handle:
        writer.write(handle)

    validator = DeliveryValidator(SimpleNamespace())
    report = validator.validate_structure(docx, pdf, expected_sections=["第一章", "第二章"])
    codes = {item["code"] for item in report["findings"]}
    assert "D5_EMPTY_SECTION" in codes
    assert "D5_PLACEHOLDER_WORD" in codes
    assert "D5_UNRENDERED_DIRECTIVE" in codes
    assert "D5_INTERNAL_PATH" in codes
