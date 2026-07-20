from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace

from docx import Document
from pypdf import PdfWriter

from app.post_export_acceptance import PostExportQualityLifecycleManager
from app.post_export_validator import PostExportDeliveryValidator


def _blank_pdf(path: Path) -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=595, height=842)
    with path.open("wb") as handle:
        writer.write(handle)


def test_content_issue_routes_to_responsible_writer():
    manager = PostExportQualityLifecycleManager(None)
    route = manager.route_delivery_finding(
        {
            "code": "D5_CONTENT_PLACEHOLDER_WORD",
            "category": "CONTENT",
            "target_type": "SECTION_CANDIDATE",
            "responsible_section_ids": ["section-method"],
        }
    )
    assert route.owner == "WRITING_AGENT"
    assert route.owner_kind == "AGENT"
    assert route.reviewer_prompt_id == "P-INTEGRATION-CRITIC"


def test_layout_issue_routes_to_export_engineering():
    manager = PostExportQualityLifecycleManager(None)
    route = manager.route_delivery_finding(
        {
            "code": "D6_TEXT_OVERLAP_RISK",
            "category": "FORMAT",
            "target_type": "PDF_LAYOUT",
        }
    )
    assert route.owner == "EXPORT_ENGINEERING"
    assert route.owner_kind == "ENGINEERING"
    assert route.stage_prompt_ids == ("EXPORT_ENGINEERING",)
    assert route.reviewer_prompt_id == "DELIVERY_VALIDATOR"


def test_reviewed_content_loss_is_detected_and_scoped(tmp_path: Path):
    docx = tmp_path / "delivery.docx"
    document = Document()
    document.add_heading("技术路线", level=1)
    document.add_paragraph("只保留了部分正文。")
    document.save(docx)
    pdf = tmp_path / "delivery.pdf"
    _blank_pdf(pdf)

    validator = PostExportDeliveryValidator(SimpleNamespace())
    report = validator.validate_structure(
        docx,
        pdf,
        expected_sections=["技术路线"],
        expected_candidates=[
            {
                "section_id": "section-method",
                "section_title": "技术路线",
                "candidate_id": "candidate-1",
                "paragraphs": ["只保留了部分正文。", "缺失的已审查方法与验证内容。"],
            }
        ],
    )
    findings = {item["code"]: item for item in report["findings"]}
    assert "D5_REVIEWED_CONTENT_LOSS" in findings
    assert findings["D5_REVIEWED_CONTENT_LOSS"]["owner"] == "EXPORT_ENGINEERING"
    assert findings["D5_REVIEWED_CONTENT_LOSS"]["responsible_section_ids"] == ["section-method"]


def test_visible_manifest_counts_structured_blocks():
    validator = PostExportDeliveryValidator(SimpleNamespace())
    manifest = validator._expected_visible_manifest(
        [
            {
                "section_id": "section-method",
                "paragraphs": [
                    "[[H2]]方法步骤",
                    "[[TABLE]]对象|指标\n算法|误差",
                    "[[FORMULA]]J=Q+T+S",
                    "正文结论。",
                ],
            }
        ]
    )
    assert manifest["table_count"] == 1
    assert manifest["figure_count"] == 0
    assert {item["kind"] for item in manifest["units"]} >= {
        "HEADING",
        "TABLE_CELL",
        "FORMULA",
        "PARAGRAPH",
    }


def test_pdf_linear_text_parity_does_not_treat_table_cells_as_contiguous_prose(tmp_path: Path):
    docx = tmp_path / "table-delivery.docx"
    document = Document()
    document.add_heading("研究问题", level=1)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "研究问题"
    table.cell(0, 1).text = "核心输出"
    table.cell(1, 0).text = "恢复跨制品可追踪关系"
    table.cell(1, 1).text = "异构关联图与证据路径"
    document.save(docx)
    pdf = tmp_path / "table-delivery.pdf"
    _blank_pdf(pdf)

    validator = PostExportDeliveryValidator(SimpleNamespace())
    expected = validator._expected_visible_manifest([
        {
            "section_id": "section-objective",
            "section_title": "研究问题",
            "candidate_id": "candidate-table",
            "paragraphs": [
                "[[TABLE]]研究问题|核心输出\n恢复跨制品可追踪关系|异构关联图与证据路径"
            ],
        }
    ])
    assert all(item["kind"] == "TABLE_CELL" for item in expected["units"])
    # Table cells are covered by structural and visual checks rather than an
    # order-sensitive linear PDF text comparison.
    pdf_parity_units = [item for item in expected["units"] if item.get("kind") != "TABLE_CELL"]
    assert pdf_parity_units == []


def _acceptance_db(tmp_path: Path):
    import json

    from app.db import Database
    from app.util import utc_now

    db = Database(tmp_path / "post-export.sqlite3")
    now = utc_now()
    db.execute(
        "INSERT INTO projects(id,name,description,security_level,config_json,created_at,updated_at) VALUES(?,?,?,?,?,?,?)",
        ("project-1", "Proposal", "validator revalidation", "PUBLIC", "{}", now, now),
    )
    db.execute(
        "INSERT INTO workflows(id,project_id,workflow_type,status,current_step,state_json,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?)",
        ("wf-4", "project-1", "WF-4_PROPOSAL_AUTHORING", "COMPLETED", 99, "{}", now, now),
    )
    output = {
        "result": {
            "scheme_profile": {
                "rules": [
                    {
                        "rule_id": "rule-pages",
                        "rule_type": "PAGE_OR_WORD_LIMIT",
                        "mandatory": True,
                        "statement": "正文16—20页，参考文献页不计入正文页数。",
                    }
                ]
            }
        }
    }
    db.execute(
        "INSERT INTO prompt_runs(id,project_id,workflow_id,prompt_id,status,model_id,endpoint_id,input_hash,output_hash,input_json,output_json,error,duration_ms,created_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
        (
            "scheme-run", "project-1", "wf-4", "P-SCHEME-EXTRACT", "PASS",
            "model", "endpoint", "input", "output", "{}",
            json.dumps(output, ensure_ascii=False), None, 1, now,
        ),
    )
    return db


def _legacy_page_finding():
    return {
        "code": "D5_GUIDE_PAGE_COUNT_OUT_OF_RANGE",
        "severity": "P1",
        "message": "交付PDF页数不满足申请指南约束：要求16—20页，实际21页。",
        "location": "PDF pages",
        "blocking": True,
        "category": "CONTENT",
        "target_type": "FULL_PROPOSAL_CANDIDATE_SET",
        "evidence": {
            "constraint": {"min": 16, "max": 20},
            "actual": 21,
            "source_rule_ids": ["rule-pages"],
        },
    }


def test_validator_revision_is_recorded_in_structure_report(tmp_path: Path):
    docx = tmp_path / "revision.docx"
    document = Document()
    document.add_heading("摘要", level=1)
    document.add_paragraph("这是用于验证版本记录的正文内容，长度足以被正文清单识别。")
    document.save(docx)
    pdf = tmp_path / "revision.pdf"
    _blank_pdf(pdf)
    validator = PostExportDeliveryValidator(SimpleNamespace())
    report = validator.validate_structure(
        docx,
        pdf,
        expected_sections=["摘要"],
        expected_candidates=[
            {
                "section_id": "S01",
                "section_title": "摘要",
                "paragraphs": ["这是用于验证版本记录的正文内容，长度足以被正文清单识别。"],
            }
        ],
    )
    assert report["validator_revision"] == validator.VALIDATOR_REVISION


def test_same_validator_revision_cannot_reopen_same_content_candidate(tmp_path: Path):
    from app.post_export_acceptance import PostExportAcceptanceManager

    db = _acceptance_db(tmp_path)
    lifecycle = PostExportQualityLifecycleManager(db)
    records = lifecycle.ingest_delivery_findings(
        project_id="project-1",
        workflow_id="wf-4",
        validation_run_id="old-validation",
        findings=[_legacy_page_finding()],
    )
    manager = PostExportAcceptanceManager(
        db, SimpleNamespace(), exporter=SimpleNamespace(delivery_validator=PostExportDeliveryValidator(SimpleNamespace()))
    )
    previous = {
        "status": "REVISE_CONTENT",
        "validator_revision": PostExportDeliveryValidator.VALIDATOR_REVISION,
        "candidate_snapshot": {"candidate_set_hash": "same"},
        "finding_ids": [records[0]["finding_id"]],
    }
    assert manager._prepare_validator_revalidation(
        project_id="project-1",
        workflow_id="wf-4",
        previous_attempt=previous,
        current_validator_revision=PostExportDeliveryValidator.VALIDATOR_REVISION,
    ) is None
    latest = lifecycle.list_findings("project-1", workflow_id="wf-4")[0]
    assert latest["responsibility"]["owner"] == "WRITING_AGENT"


def test_legacy_total_page_false_positive_can_be_reclassified_and_independently_verified(tmp_path: Path):
    from app.post_export_acceptance import PostExportAcceptanceManager
    from app.util import sha256_json

    db = _acceptance_db(tmp_path)
    lifecycle = PostExportQualityLifecycleManager(db)
    records = lifecycle.ingest_delivery_findings(
        project_id="project-1",
        workflow_id="wf-4",
        validation_run_id="old-validation",
        findings=[_legacy_page_finding()],
    )
    finding_id = records[0]["finding_id"]
    manager = PostExportAcceptanceManager(
        db, SimpleNamespace(), exporter=SimpleNamespace(delivery_validator=PostExportDeliveryValidator(SimpleNamespace()))
    )
    plan = manager._prepare_validator_revalidation(
        project_id="project-1",
        workflow_id="wf-4",
        previous_attempt={
            "status": "REVISE_CONTENT",
            "candidate_snapshot": {"candidate_set_hash": "same"},
            "finding_ids": [finding_id],
        },
        current_validator_revision=PostExportDeliveryValidator.VALIDATOR_REVISION,
    )
    assert plan is not None
    assert plan["previous_validator_revision"] == "legacy-unversioned"
    assert plan["reclassified_finding_ids"] == [finding_id]
    reclassified = lifecycle.list_findings("project-1", workflow_id="wf-4")[0]
    assert reclassified["responsibility"]["owner"] == "DELIVERY_VALIDATOR_ENGINEERING"
    assert reclassified["lifecycle"]["reclassification_evidence"]

    lifecycle.add_repair_evidence(
        finding_id,
        project_id="project-1",
        prompt_id="DELIVERY_VALIDATOR_ENGINEERING",
        run_id="validator-fix",
    )
    lifecycle.verify_finding(
        finding_id,
        project_id="project-1",
        reviewer="DELIVERY_VALIDATOR",
        review_run_id="new-validation",
        review_hash=sha256_json({"absent_finding_id": finding_id}),
    )
    verified = lifecycle.list_findings("project-1", workflow_id="wf-4")[0]
    assert verified["lifecycle"]["state"] == "VERIFIED"


def test_unrelated_content_finding_cannot_use_validator_revision_revalidation(tmp_path: Path):
    from app.post_export_acceptance import PostExportAcceptanceManager

    db = _acceptance_db(tmp_path)
    lifecycle = PostExportQualityLifecycleManager(db)
    unrelated = _legacy_page_finding()
    unrelated.update(
        {
            "code": "D5_GUIDE_REFERENCE_COUNT_OUT_OF_RANGE",
            "message": "参考文献数量不足。",
            "location": "DOCX references",
            "evidence": {"constraint": {"min": 30, "max": 40}, "actual": 20},
        }
    )
    records = lifecycle.ingest_delivery_findings(
        project_id="project-1",
        workflow_id="wf-4",
        validation_run_id="old-validation",
        findings=[unrelated],
    )
    manager = PostExportAcceptanceManager(
        db, SimpleNamespace(), exporter=SimpleNamespace(delivery_validator=PostExportDeliveryValidator(SimpleNamespace()))
    )
    assert manager._prepare_validator_revalidation(
        project_id="project-1",
        workflow_id="wf-4",
        previous_attempt={
            "status": "REVISE_CONTENT",
            "candidate_snapshot": {"candidate_set_hash": "same"},
            "finding_ids": [records[0]["finding_id"]],
        },
        current_validator_revision=PostExportDeliveryValidator.VALIDATOR_REVISION,
    ) is None
