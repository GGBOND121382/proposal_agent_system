from __future__ import annotations

import json
import zipfile
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document
from PIL import Image

from app.db import Database
from app.figure_protocol import FigureDirective, artifact_reference
from app.research_mermaid_export import (
    ResearchMermaidExportError,
    ResearchMermaidExportPipeline,
)
from app.skills.base import SkillResult
from app.util import sha256_bytes, sha256_text, write_json


class _FakeSkillExecutor:
    def __init__(self, data_dir: Path):
        self.data_dir = data_dir

    def execute(self, skill_id, payload, *, project_id, workflow_id, security_level):
        if skill_id == "public_research.archive":
            return self._research(project_id, workflow_id)
        if skill_id == "mermaid.render":
            return self._mermaid(payload, project_id, workflow_id)
        raise AssertionError(skill_id)

    def _research(self, project_id: str, workflow_id: str | None) -> SkillResult:
        root = self.data_dir / "research_archive" / project_id / "fixed-session"
        raw_dir = root / "raw"
        text_dir = root / "text"
        meta_dir = root / "metadata"
        connector_dir = root / "connector"
        for directory in (raw_dir, text_dir, meta_dir, connector_dir):
            directory.mkdir(parents=True, exist_ok=True)
        connector = connector_dir / "connector_response.json"
        connector.write_text('{"connector":"fixture"}', encoding="utf-8")

        records = []
        sources = []
        for source_id, title, url, text, rank in (
            (
                "src-official",
                "Official Evaluation Framework",
                "https://example.gov/framework",
                "The official framework retains reproducible evidence, baseline comparisons and documented limitations.",
                95,
            ),
            (
                "src-review",
                "Recent Benchmark Review",
                "https://example.edu/review",
                "This 2025 review compares recent benchmark baselines and explains limitations and open challenges.",
                85,
            ),
        ):
            raw = json.dumps({"title": title, "url": url, "content_text": text}, ensure_ascii=False, indent=2).encode("utf-8")
            raw_path = raw_dir / f"{source_id}.json"
            text_path = text_dir / f"{source_id}.txt"
            meta_path = meta_dir / f"{source_id}.json"
            raw_path.write_bytes(raw)
            text_path.write_text(text, encoding="utf-8")
            record = {
                "source_id": source_id,
                "title": title,
                "url": url,
                "final_url": url,
                "published_at": "2025-01-01",
                "publisher": "Public Authority" if rank == 95 else "University",
                "source_type": "GOVERNMENT" if rank == 95 else "PEER_REVIEWED_PAPER",
                "authority_rank": rank,
                "snapshot_sha256": sha256_bytes(raw),
                "text_sha256": sha256_text(text),
                "raw_path": str(raw_path),
                "text_path": str(text_path),
                "metadata_path": str(meta_path),
                "excerpt": text,
            }
            write_json(meta_path, record)
            records.append(record)
            sources.append(
                {
                    "source_id": source_id,
                    "source_type": "PUBLIC_SOURCE",
                    "document_version_id": None,
                    "section_id": None,
                    "span_start": None,
                    "span_end": None,
                    "quoted_text": f"{title} | {url}",
                    "source_hash": record["snapshot_sha256"],
                    "authority_rank": rank,
                    "security_level": "PUBLIC",
                }
            )
        manifest = {
            "schema_version": "2.0",
            "project_id": project_id,
            "workflow_id": workflow_id,
            "records": records,
            "connector_response": str(connector),
        }
        manifest_path = root / "manifest.json"
        write_json(manifest_path, manifest)
        output = {
            "mode": "RECORDED_CONNECTOR_INTEGRATION",
            "archive_root": str(root),
            "archive_manifest": str(manifest_path),
            "archive_verification": {"status": "PASS"},
            "plan_validation": {"status": "PASS"},
            "coverage": {
                "status": "PASS",
                "dimensions": {
                    "recent_work": {"status": "PASS", "source_ids": ["src-review"]},
                    "comparable_baselines": {"status": "PASS", "source_ids": ["src-review"]},
                    "limitation_mechanisms": {"status": "PASS", "source_ids": ["src-review"]},
                },
                "uncovered_queries": [],
            },
            "issues": [],
            "sources": sources,
            "source_catalog": records,
        }
        return SkillResult("PASS", output, [], [str(manifest_path)])

    def _mermaid(self, payload: dict, project_id: str, workflow_id: str | None) -> SkillResult:
        section_id = payload["section_id"]
        root = self.data_dir / "diagram_artifacts" / project_id / section_id
        root.mkdir(parents=True, exist_ok=True)
        mmd = root / f"{section_id}.mmd"
        svg = root / f"{section_id}.svg"
        png = root / f"{section_id}.png"
        metadata = root / f"{section_id}.json"
        source = payload["mermaid_source"]
        mmd.write_text(source, encoding="utf-8")
        svg.write_text("<svg xmlns='http://www.w3.org/2000/svg'><text>ok</text></svg>", encoding="utf-8")
        Image.new("RGB", (900, 500), "white").save(png)
        output = {
            "section_id": section_id,
            "caption": payload["caption"],
            "source_reference": artifact_reference(mmd, self.data_dir),
            "svg_reference": artifact_reference(svg, self.data_dir),
            "png_reference": artifact_reference(png, self.data_dir),
            "source_sha256": sha256_text(source),
            "svg_sha256": sha256_text(svg.read_text(encoding="utf-8")),
            "png_sha256": sha256_bytes(png.read_bytes()),
            "cache_hit": False,
        }
        output["figure_marker"] = FigureDirective(
            output["png_reference"],
            payload["caption"],
            float(payload.get("width_cm") or 15.0),
            output["source_reference"],
        ).marker()
        write_json(metadata, output)
        return SkillResult(
            "PASS",
            output,
            [],
            [str(mmd), str(svg), str(png), str(metadata)],
        )


class _FakeExporter:
    def __init__(self, data_dir: Path):
        self.data_dir = data_dir
        self.candidates: list[dict] = []

    def _candidate_runs(self, project_id: str):
        return self.candidates

    def export(self, project_id: str) -> Path:
        path = self.data_dir / "exports" / "s3.docx"
        path.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        document.add_heading("S3", level=1)
        document.add_paragraph("validated")
        document.save(path)
        path.with_suffix(".integrity.json").write_text("{}", encoding="utf-8")
        path.with_suffix(".manifest.json").write_text("{}", encoding="utf-8")
        return path

    def export_package(self, project_id: str, document_path: Path) -> Path:
        pdf = document_path.with_suffix(".pdf")
        pdf.write_bytes(b"%PDF-1.4\n% unit-test\n")
        validation = {
            "status": "PASS",
            "finding_count": 0,
            "blocking_finding_count": 0,
        }
        write_json(document_path.with_suffix(".delivery-validation.json"), validation)
        package = document_path.with_suffix(".zip")
        with zipfile.ZipFile(package, "w") as archive:
            archive.write(document_path, arcname=document_path.name)
            archive.write(pdf, arcname=pdf.name)
        return package


def _settings(tmp_path: Path) -> SimpleNamespace:
    return SimpleNamespace(data_dir=tmp_path)


def _synthesis(source_id: str = "src-official") -> dict:
    source_hash = None
    # Hashes are supplied by the fixture executor; the known raw JSON is deterministic.
    title = "Official Evaluation Framework"
    url = "https://example.gov/framework"
    text = "The official framework retains reproducible evidence, baseline comparisons and documented limitations."
    raw = json.dumps({"title": title, "url": url, "content_text": text}, ensure_ascii=False, indent=2).encode("utf-8")
    if source_id == "src-official":
        source_hash = sha256_bytes(raw)
    else:
        source_hash = "0" * 64
    return {
        "claims": [
            {
                "claim_id": "claim-1",
                "claim_text": "公开框架要求保留可复核证据。",
                "claim_type": "PUBLIC_CLAIM",
                "subject_id": None,
                "temporal_status": "CURRENT",
                "qualifiers": [],
                "numeric_values": [],
                "source_refs": [
                    {
                        "source_id": source_id,
                        "source_type": "PUBLIC_SOURCE",
                        "quoted_text": title,
                        "source_hash": source_hash,
                    }
                ],
                "knowledge_status": "DOCUMENT_EXTRACTED",
                "security_level": "PUBLIC",
            }
        ],
        "source_comparisons": [],
        "conflicts": [],
        "limitations": [],
        "coverage_summary": "覆盖完整",
    }


def _diagrams(source_ids=None) -> list[dict]:
    return [
        {
            "section_id": "research-evidence-flow",
            "caption": "公开研究证据链",
            "width_cm": 13.5,
            "mermaid_source": "flowchart LR\n A[Research] --> B[Claim] --> C[Export]",
            "claim_ids": ["claim-1"],
            "source_ids": source_ids or ["src-official"],
        }
    ]


def _pipeline(tmp_path: Path):
    db = Database(tmp_path / "db.sqlite3")
    executor = _FakeSkillExecutor(tmp_path)
    exporter = _FakeExporter(tmp_path)
    pipeline = ResearchMermaidExportPipeline(
        db,
        _settings(tmp_path),
        skill_executor=executor,
        exporter=exporter,
    )
    return pipeline, exporter


def _prepare(pipeline: ResearchMermaidExportPipeline, diagrams=None):
    return pipeline.prepare(
        project_id="project-s3",
        workflow_id="workflow-s3",
        security_level="PUBLIC",
        research_plan={"queries": ["official reproducible evidence benchmark limitations"]},
        research_request={"provider": "connector", "connector_file": "fixture.json"},
        synthesis=_synthesis(),
        diagrams=diagrams or _diagrams(),
        acceptance_mode="RECORDED_CONNECTOR_INTEGRATION",
    )


def test_prepare_writes_portable_restart_verified_checkpoint(tmp_path: Path):
    pipeline, _ = _pipeline(tmp_path)
    prepared = _prepare(pipeline)
    assert prepared["status"] == "WAITING_FOR_EXPRESSION_APPROVED_CONTENT"
    assert prepared["claim_binding"]["status"] == "PASS"
    assert prepared["verification"]["reference"].startswith("artifact://")
    checkpoint_path = pipeline._resolve_reference(prepared["checkpoint"]["reference"])
    checkpoint_text = checkpoint_path.read_text(encoding="utf-8")
    assert str(tmp_path.resolve()) not in checkpoint_text
    assert pipeline.verify_prepared(checkpoint_path)["status"] == "PASS"


def test_prepare_blocks_diagram_source_not_bound_to_claim(tmp_path: Path):
    pipeline, _ = _pipeline(tmp_path)
    with pytest.raises(ResearchMermaidExportError, match="not bound"):
        _prepare(pipeline, diagrams=_diagrams(["src-review"]))


def test_finalize_requires_each_marker_once_in_approved_content(tmp_path: Path):
    pipeline, exporter = _pipeline(tmp_path)
    prepared = _prepare(pipeline)
    exporter.candidates = [
        {
            "section_id": "s1",
            "section_title": "公开研究",
            "candidate_id": "candidate-1",
            "run_id": "polish-1",
            "expression_critic_run_id": "critic-1",
            "paragraphs": ["正文中没有图。"],
        }
    ]
    with pytest.raises(ResearchMermaidExportError, match="exactly once"):
        pipeline.finalize(project_id="project-s3", checkpoint=prepared)

def test_finalize_exports_and_bundles_verified_chain(tmp_path: Path):
    pipeline, exporter = _pipeline(tmp_path)
    prepared = _prepare(pipeline)
    exporter.candidates = [
        {
            "section_id": "s1",
            "section_title": "公开研究",
            "candidate_id": "candidate-1",
            "run_id": "polish-1",
            "expression_critic_run_id": "critic-1",
            "paragraphs": ["公开研究正文。", prepared["required_figure_markers"][0]],
        }
    ]
    result = pipeline.finalize(project_id="project-s3", checkpoint=prepared)
    assert result["status"] == "PASS"
    assert result["delivery"]["validation"]["reference"].startswith("artifact://")
    bundle = pipeline._resolve_reference(result["evidence_bundle"]["reference"])
    assert bundle.is_file()
    with zipfile.ZipFile(bundle) as archive:
        names = set(archive.namelist())
    assert any(name.endswith("research-manifest.portable.json") for name in names)
    assert any(name.endswith("raw/src-official.json") for name in names)
    assert any(name.endswith("text/src-official.txt") for name in names)
    assert any(name.endswith("metadata/src-official.json") for name in names)
    assert any(name.endswith("connector/connector_response.json") for name in names)
    assert any(name.endswith("s3.zip") for name in names)


def test_restart_verification_detects_tampered_mermaid_artifact(tmp_path: Path):
    pipeline, _ = _pipeline(tmp_path)
    prepared = _prepare(pipeline)
    png_record = next(
        record
        for record in prepared["diagrams"][0]["artifacts"]
        if record["artifact_type"] == "MERMAID_PNG"
    )
    pipeline._resolve_reference(png_record["reference"]).write_bytes(b"tampered")
    verification = pipeline.verify_prepared(prepared)
    assert verification["status"] == "FAIL"
    assert "S3_ARTIFACT_HASH_MISMATCH" in {
        item["code"] for item in verification["findings"]
    }


def test_final_restart_verification_rechecks_research_artifacts(tmp_path: Path):
    pipeline, exporter = _pipeline(tmp_path)
    prepared = _prepare(pipeline)
    exporter.candidates = [
        {
            "section_id": "s1",
            "section_title": "公开研究",
            "candidate_id": "candidate-1",
            "run_id": "polish-1",
            "expression_critic_run_id": "critic-1",
            "paragraphs": [prepared["required_figure_markers"][0]],
        }
    ]
    result = pipeline.finalize(project_id="project-s3", checkpoint=prepared)
    portable = pipeline._resolve_reference(
        prepared["research"]["portable_manifest"]["reference"]
    )
    manifest = json.loads(portable.read_text(encoding="utf-8"))
    pipeline._resolve_reference(manifest["records"][0]["raw_path"]).write_bytes(b"tampered")
    verification = pipeline.verify_final(result)
    assert verification["status"] == "FAIL"
    assert "S3_FINAL_S3_RESEARCH_ARTIFACT_HASH_MISMATCH" in {
        item["code"] for item in verification["findings"]
    }
