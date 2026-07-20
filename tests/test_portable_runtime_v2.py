from __future__ import annotations

import pytest

import asyncio
import json
from pathlib import Path
from types import SimpleNamespace

from docx import Document
from pypdf import PdfWriter

from app.context_base import ContextBuilder
from app.delivery_validator import DeliveryValidator
from app.g3_runtime_executor import G3RuntimePromptExecutor
from app.human_gate_bridge import FileHumanGateBridge
from app.post_export_validator import PostExportDeliveryValidator
from app.proposal_constraints import extract_hard_constraints, merge_contract_constraints
from app.runtime_gateway_factory import PortableModelGateway
from app.task_instruction import instruction_text, intended_uses


def test_task_instruction_object_normalizes_to_scalar_and_string_list():
    value = {
        "objective": "生成完整申请书",
        "constraints": ["不得虚构成果", "保持UNKNOWN"],
        "deliverables": ["DOCX", "PDF"],
    }
    assert instruction_text(value) == "生成完整申请书"
    assert intended_uses(value) == [
        "生成完整申请书",
        "不得虚构成果",
        "保持UNKNOWN",
        "DOCX",
        "PDF",
    ]
    assert all(isinstance(item, str) for item in intended_uses(value))


def test_gateway_factory_requires_explicit_bridge_mode(monkeypatch, tmp_path: Path):
    import app.runtime_gateway_factory as module

    class FakeHttp:
        def __init__(self, settings, pack):
            self.kind = "http"

    class FakeBridge:
        def __init__(self, settings, pack):
            self.kind = "bridge"

    monkeypatch.setattr(module, "G3AuditedModelGateway", FakeHttp)
    monkeypatch.setattr(module, "ChatBridgeModelGateway", FakeBridge)
    settings = SimpleNamespace(model_gateway_mode="OPENAI_COMPATIBLE", chat_bridge_dir=tmp_path)
    assert PortableModelGateway(settings, object()).kind == "http"
    settings.model_gateway_mode = "CHAT_BRIDGE"
    assert PortableModelGateway(settings, object()).kind == "bridge"
    settings.chat_bridge_dir = None
    monkeypatch.delenv("CHAT_BRIDGE_DIR", raising=False)
    try:
        PortableModelGateway(settings, object())
    except ValueError as exc:
        assert "CHAT_BRIDGE_DIR" in str(exc)
    else:
        raise AssertionError("bridge mode must require a directory")


def test_human_gate_bridge_binds_context_hash_and_role(tmp_path: Path):
    gate = {
        "id": "gate-1",
        "project_id": "project-1",
        "workflow_id": "wf-1",
        "gate_type": "SCHEME_CONFIRMATION",
        "target_id": "artifact-1",
        "required_role": "PROJECT_OWNER",
        "allowed_actions": ["CONFIRM", "RETURN"],
        "questions": [{"question_id": "q1", "question": "确认？"}],
        "context_hash": "a" * 64,
    }

    class Engine:
        def __init__(self):
            self.calls = []

        def decide_gate(self, gate_id, **kwargs):
            self.calls.append((gate_id, kwargs))
            return {"id": gate_id, "status": "APPROVED"}

    bridge = FileHumanGateBridge(tmp_path, poll_seconds=0.01, timeout_seconds=2)
    request_path = bridge.publish(gate)
    assert json.loads(request_path.read_text(encoding="utf-8"))["context_hash"] == "a" * 64
    bridge.responses_dir.mkdir(parents=True, exist_ok=True)
    (bridge.responses_dir / "gate-1.json").write_text(
        json.dumps(
            {
                "gate_id": "gate-1",
                "context_hash": "a" * 64,
                "action": "CONFIRM",
                "decided_by": "gpt",
                "decided_role": "PROJECT_OWNER",
                "answers": [],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    engine = Engine()
    result = asyncio.run(bridge.wait_and_apply(engine, gate))
    assert result["status"] == "APPROVED"
    assert engine.calls[0][1]["context_hash"] == "a" * 64
    assert (bridge.consumed_dir / "gate-1.json").is_file()


def test_live_compaction_preserves_structural_arrays():
    value = {
        "linked_sections": [{"section_id": f"sec-{i}"} for i in range(14)],
        "mandatory_sections": [f"sec-{i}" for i in range(14)],
        "tasks": [{"revision_task_id": f"task-{i}"} for i in range(14)],
        "items": [{"item_id": f"item-{i}"} for i in range(34)],
        "relations": [{"relation_id": f"rel-{i}"} for i in range(19)],
    }
    compact = G3RuntimePromptExecutor._compact_live_value(value, aggressive=True)
    assert len(compact["linked_sections"]) == 14
    assert len(compact["mandatory_sections"]) == 14
    assert len(compact["tasks"]) == 14
    assert len(compact["items"]) == 34
    assert len(compact["relations"]) == 19


def test_scoped_plan_uses_contract_order_not_first_task_fallback():
    contracts = [
        {"section_id": f"sec-{i}", "title": f"章节{i}", "argument_function": f"职责{i}"}
        for i in range(14)
    ]
    tasks = [
        {
            "revision_task_id": f"task-{i}",
            "objective": f"其他表述{i}",
            "required_input_ids": [],
            "issue_ids": [f"issue-{i}"],
            "acceptance_rules": ["完成本章"],
        }
        for i in range(14)
    ]
    plan = {"tasks": tasks, "dependencies": []}
    architecture = {"section_contracts": contracts}
    scoped = ContextBuilder._scoped_plan(plan, architecture, contracts[11])
    assert [item["revision_task_id"] for item in scoped["tasks"]] == ["task-11"]




def test_scheme_constraints_parse_natural_chinese_figure_table_minima():
    scheme = {
        "rules": [
            {
                "rule_id": "rule-figures", "mandatory": True,
                "rule_type": "DOCUMENT_STRUCTURE",
                "statement": "正文至少包含5幅有效图形，且均须在正文中明确引用并服务于论证。",
            },
            {
                "rule_id": "rule-tables", "mandatory": True,
                "rule_type": "DOCUMENT_STRUCTURE",
                "statement": "正文至少包含6张有效表格，且均须在正文中明确引用并服务于论证。",
            },
        ]
    }
    constraints = extract_hard_constraints(scheme)
    assert constraints["minimum_figures"] == 5
    assert constraints["minimum_tables"] == 6
    assert constraints["source_rule_ids"] == ["rule-figures", "rule-tables"]


def test_scheme_constraints_are_normalized_and_merged():
    scheme = {
        "rules": [
            {"rule_id": "r-page", "mandatory": True, "rule_type": "PAGE_OR_WORD_LIMIT", "statement": "正文16—20页"},
            {"rule_id": "r-ref", "mandatory": True, "rule_type": "DOCUMENT_STRUCTURE", "statement": "参考文献30—40篇"},
            {"rule_id": "r-visual", "mandatory": True, "rule_type": "DOCUMENT_STRUCTURE", "statement": "至少5图6表"},
        ]
    }
    constraints = extract_hard_constraints(scheme)
    assert constraints["main_body_pages"] == {"min": 16, "max": 20}
    assert constraints["references"] == {"min": 30, "max": 40}
    assert constraints["minimum_figures"] == 5
    assert constraints["minimum_tables"] == 6
    merged = merge_contract_constraints(
        {
            "contract_id": "c1",
            "document_type": "RESEARCH_PROPOSAL",
            "funding_scheme": None,
            "primary_evaluation_logic": "SCIENTIFIC_MERIT",
            "target_evaluators": [],
            "max_main_pages": None,
            "max_core_research_questions": 4,
            "mandatory_sections": [],
            "appendix_only_topics": [],
            "forbidden_main_body_topics": [],
            "status": "CONFIRMED",
        },
        constraints,
    )
    assert merged["min_main_pages"] == 16
    assert merged["max_main_pages"] == 20
    assert merged["min_reference_count"] == 30
    assert merged["min_figure_count"] == 5


def test_placeholder_detector_distinguishes_boundary_statement():
    pattern = DeliveryValidator.PLACEHOLDER_PATTERNS["PLACEHOLDER_WORD"]
    assert not pattern.search("对于尚未获得的数据，设置待补充条件并保持UNKNOWN状态。")
    assert pattern.search("待补充：申请人论文清单")
    assert pattern.search("TODO")


def test_delivery_validator_enforces_guide_counts(tmp_path: Path):
    docx = tmp_path / "proposal.docx"
    pdf = tmp_path / "proposal.pdf"
    document = Document()
    document.add_heading("摘要", level=1)
    document.add_paragraph("本项目研究需求变更影响分析与智能测试推荐。")
    document.save(docx)
    writer = PdfWriter()
    writer.add_blank_page(width=595, height=842)
    with pdf.open("wb") as handle:
        writer.write(handle)

    validator = PostExportDeliveryValidator(SimpleNamespace())
    report = validator.validate_structure(
        docx,
        pdf,
        expected_sections=["摘要"],
        expected_candidates=[
            {
                "section_id": "sec-abstract",
                "section_title": "摘要",
                "candidate_id": "candidate-1",
                "paragraphs": ["本项目研究需求变更影响分析与智能测试推荐。"],
            }
        ],
        expected_constraints={
            "source_rule_ids": ["r-page", "r-ref", "r-visual"],
            "main_body_pages": {"min": 16, "max": 20},
            "references": {"min": 30, "max": 40},
            "minimum_figures": 5,
            "minimum_tables": 6,
        },
    )
    codes = {item["code"] for item in report["findings"]}
    assert "D5_GUIDE_PAGE_COUNT_OUT_OF_RANGE" in codes
    assert "D5_GUIDE_REFERENCE_COUNT_OUT_OF_RANGE" in codes
    assert "D5_GUIDE_FIGURE_MINIMUM_UNMET" in codes
    assert "D5_GUIDE_TABLE_MINIMUM_UNMET" in codes


def test_repair_allowlist_accepts_named_selector_but_keeps_explicit_index_strict():
    from app.track_b import TrackBAgentPromptValidator

    validator = object.__new__(TrackBAgentPromptValidator)
    payload = {
        "allowed_paths": ["content.section_contracts.word_budget"],
        "protected_paths": [],
        "findings_to_repair": [{"code": "QG_BUDGET"}],
    }
    result = {
        "changed_paths": ["content.section_contracts[LITERATURE_REVIEW].word_budget"],
        "resolved_finding_codes": ["QG_BUDGET"],
    }
    assert validator._audit_repair_scope(payload, result) == []
    strict_payload = {**payload, "allowed_paths": ["content.section_contracts[3].word_budget"]}
    findings = validator._audit_repair_scope(
        strict_payload,
        {**result, "changed_paths": ["content.section_contracts[4].word_budget"]},
    )
    assert any(item.code == "QG_REPAIR_PATH_OUTSIDE_ALLOWLIST" for item in findings)


def test_markdown_skeleton_preserves_empty_explicit_headings():
    from app.documents import parse_document

    parsed = parse_document(
        "skeleton.md",
        "# 项目名称\n\n# 摘要\n\n# 研究内容\n\n# 风险与结论\n".encode("utf-8"),
        "CURRENT_PROPOSAL",
        "PUBLIC",
    )
    titles = [item["title"] for item in parsed["sections"] if item["title"] != "全文"]
    assert titles == ["项目名称", "摘要", "研究内容", "风险与结论"]
    assert all(item["text"] == "" for item in parsed["sections"] if item["title"] != "全文")


def test_portable_driver_scopes_gates_to_current_workflow_and_children():
    from scripts.run_portable_workflow import _open_gates

    class Engine:
        def __init__(self):
            self.calls = []

        def get(self, workflow_id):
            assert workflow_id == "wf-current"
            return {
                "id": workflow_id,
                "state": {
                    "full_proposal_children": {
                        "GROUP_1": {"workflow_id": "wf-child-1"},
                        "GROUP_2": {"workflow_id": "wf-child-2"},
                    }
                },
            }

        def list_gates(self, *, workflow_id=None, project_id=None):
            self.calls.append((workflow_id, project_id))
            assert project_id is None
            return [
                {
                    "id": f"gate-{workflow_id}",
                    "workflow_id": workflow_id,
                    "status": "OPEN",
                    "created_at": "2026-07-17T00:00:00Z",
                }
            ]

    engine = Engine()
    gates = _open_gates(engine, "wf-current")
    assert [item["workflow_id"] for item in gates] == [
        "wf-current",
        "wf-child-1",
        "wf-child-2",
    ]
    assert engine.calls == [
        ("wf-current", None),
        ("wf-child-1", None),
        ("wf-child-2", None),
    ]


def test_workflow_gate_scope_excludes_unrelated_project_workflows():
    from app.human_gate_bridge import workflow_gate_scope_ids

    class DB:
        def fetchall(self, sql, params):
            assert params == ("project-1",)
            return [
                {"id": "wf-parent", "state_json": json.dumps({"authoring_child_workflow_ids": ["wf-child"]})},
                {"id": "wf-child", "state_json": json.dumps({"parent_workflow_id": "wf-parent"})},
                {"id": "wf-stale", "state_json": json.dumps({})},
            ]

    class Engine:
        db = DB()
        def get(self, workflow_id):
            return {"id": workflow_id, "project_id": "project-1"}

    assert workflow_gate_scope_ids(Engine(), "wf-parent") == {"wf-parent", "wf-child"}


def test_gate_decision_replay_is_idempotent(tmp_path, monkeypatch):
    from tests.test_runtime import create_project, runtime as runtime_fixture
    settings, pack, db, router, builder, executor, engine, exporter = runtime_fixture.__wrapped__(tmp_path, monkeypatch)
    project_id = create_project(db)
    workflow = engine.start(project_id, "WF-1_PROJECT_INTAKE")
    gate_id = engine._create_gate(workflow, "SCHEME_CONFIRMATION", target_id="run-1", questions=[])
    gate = engine._gate(gate_id)
    first = engine.decide_gate(
        gate_id,
        action="CONFIRM",
        decided_by="bridge",
        decided_role=gate["required_role"],
        context_hash=gate["context_hash"],
    )
    second = engine.decide_gate(
        gate_id,
        action="CONFIRM",
        decided_by="bridge-replay",
        decided_role=gate["required_role"],
        context_hash=gate["context_hash"],
    )
    assert first["status"] == second["status"] == "APPROVED"


def test_structured_task_instruction_drops_alias_fields_and_maps_values():
    structured = ContextBuilder._structured_task_instruction(
        "生成完整申请书",
        ["sec-1", "sec-2"],
        {
            "task_instruction_structured": {
                "objective": "生成完整申请书",
                "constraints": ["不得虚构成果", "正文16—20页"],
                "deliverables": ["DOCX", "PDF"],
            }
        },
    )
    assert "constraints" not in structured
    assert "deliverables" not in structured
    assert structured["specific_requirements"] == ["不得虚构成果", "正文16—20页"]
    assert structured["acceptance_preferences"] == ["DOCX", "PDF"]
    assert structured["target_section_ids"] == ["sec-1", "sec-2"]


def test_structured_task_instruction_projects_aliases_to_schema():
    raw = {
        "objective": "生成完整申请书",
        "constraints": ["正文16至20页", "不得虚构成果"],
        "deliverables": ["DOCX", "PDF"],
        "unexpected": "must not leak",
    }
    result = ContextBuilder._structured_task_instruction(
        "生成完整申请书",
        ["sec-1", "sec-2"],
        {"task_instruction": raw},
        raw_instruction=raw,
    )
    assert "constraints" not in result
    assert "deliverables" not in result
    assert "unexpected" not in result
    assert result["specific_requirements"] == ["正文16至20页", "不得虚构成果"]
    assert result["acceptance_preferences"] == ["DOCX", "PDF"]
    assert len(result["instruction_hash"]) == 64


def test_safe_package_scan_ignores_policy_labels_but_detects_outbound_identity() -> None:
    from app.context_base import ContextBuilder

    project = {"id": "project-secret-123", "name": "真实项目名称"}
    config = {
        "prohibited_external_fields": ["真实个人信息", "真实单位信息", "联系方式"],
        "prohibited_external_values": ["13800000000"],
    }
    compliant = {
        "package_id": "safe-public-001",
        "task_type": "PUBLIC_RESEARCH",
        "task_description": "检索公开的软件工程研究。",
        "queries": ["requirements traceability"],
        "allowed_context": ["公开论文与数据集"],
        "entity_placeholders": [],
        "removed_fields": ["真实个人信息", "真实单位信息", "联系方式"],
        "prohibited_outputs": ["不得输出真实个人信息"],
    }
    assert ContextBuilder._safe_package_leaked_literals(compliant, project, config) == []

    leaked_name = {**compliant, "task_description": "为真实项目名称检索公开研究"}
    assert ContextBuilder._safe_package_leaked_literals(leaked_name, project, config) == ["真实项目名称"]

    leaked_value = {**compliant, "queries": ["联系13800000000"]}
    assert ContextBuilder._safe_package_leaked_literals(leaked_value, project, config) == ["13800000000"]


def test_chat_bridge_routes_disabled_http_profile_without_weakening_security(monkeypatch) -> None:
    from types import SimpleNamespace
    from app.security import RoutingDenied, SecurityRouter

    pack = SimpleNamespace(
        endpoints={"endpoints": [{
            "endpoint_id": "online-public-primary",
            "environment": "ONLINE_PUBLIC",
            "enabled": False,
            "allowed_security_levels": ["PUBLIC"],
        }]},
        models={"models": [{
            "model_id": "online-public-primary",
            "endpoint_id": "online-public-primary",
            "provider_model_name": "",
            "enabled": False,
        }]},
        entry=lambda prompt_id: {"required_environment": "ONLINE_PUBLIC"},
        model_profile=lambda prompt_id: {"preferred_models": ["online-public-primary"], "fallback_models": []},
    )
    envelope = {"security_context": {
        "input_max_security_level": "PUBLIC",
        "online_transfer_approval_status": "APPROVED",
        "allowed_model_endpoint_ids": ["online-public-primary"],
    }}

    monkeypatch.setenv("MODEL_RUNTIME_MODE", "LIVE")
    monkeypatch.setenv("MODEL_GATEWAY_MODE", "OPENAI_COMPATIBLE")
    with pytest.raises(RoutingDenied):
        SecurityRouter(pack).route("P-PUBLIC-RESEARCH-PLAN", envelope)

    monkeypatch.setenv("MODEL_GATEWAY_MODE", "CHAT_BRIDGE")
    route = SecurityRouter(pack).route("P-PUBLIC-RESEARCH-PLAN", envelope)
    assert route.endpoint_id == "online-public-primary"
    assert route.environment == "ONLINE_PUBLIC"

    envelope["security_context"]["input_max_security_level"] = "INTERNAL"
    with pytest.raises(RoutingDenied):
        SecurityRouter(pack).route("P-PUBLIC-RESEARCH-PLAN", envelope)


def test_public_search_file_bridge_binds_request_and_covers_queries(tmp_path: Path):
    from app.public_search_bridge import FilePublicSearchBridge

    bridge = FilePublicSearchBridge(tmp_path / "search-bridge", poll_seconds=0.01, timeout_seconds=2)
    plan = {
        "plan_id": "plan-1",
        "queries": ["requirements traceability benchmark", "regression test prioritization APFD"],
    }
    request = bridge.publish(plan, 40)
    response = {
        "request_id": request["request_id"],
        "request_hash": request["request_hash"],
        "run_id": "web-run-1",
        "connector": "chatgpt-web",
        "created_at": "2026-07-17T00:00:00Z",
        "responses": [
            {
                "query": query,
                "retrieved_at": "2026-07-17T00:00:00Z",
                "results": [{
                    "title": f"Source for {query}",
                    "url": "https://example.org/source",
                    "content_text": "Verifiable public abstract and metadata content for the requested research query.",
                }],
            }
            for query in plan["queries"]
        ],
    }
    response_path = bridge.responses_dir / f"{request['request_id']}.json"
    response_path.write_text(json.dumps(response), encoding="utf-8")
    consumed = asyncio.run(bridge.request(plan, 40))
    assert consumed.is_file()
    assert json.loads(consumed.read_text())["request_hash"] == request["request_hash"]


def test_public_search_provider_recovery_requires_ready_provider():
    from app.runtime_workflows import RecoverableWorkflowEngine

    class DB:
        def __init__(self): self.events = []
        def audit(self, event, **kwargs): self.events.append(event)

    class Engine:
        _recover_status = RecoverableWorkflowEngine._recover_status
        def __init__(self, ready):
            self.db = DB()
            self.research_service = SimpleNamespace(provider_ready=lambda: ready)
            self.current = {
                "id": "wf-search", "project_id": "project-1",
                "workflow_type": "WF-3_HYBRID_ONLINE_ASSIST",
                "status": "BLOCKED", "current_step": 3,
                "state": {"last_error": "PUBLIC_SEARCH_PROVIDER is disabled"},
            }
        def get(self, workflow_id): return self.current
        def _update(self, wf, *, status=None, state=None, **kwargs):
            if status is not None: self.current["status"] = status
            if state is not None: self.current["state"] = state

    blocked = Engine(False)._recover_status(Engine(False).current)
    assert blocked["status"] == "BLOCKED"
    engine = Engine(True)
    recovered = engine._recover_status(engine.current)
    assert recovered["status"] == "RUNNING"
    assert recovered["state"]["recovered_from"] == "PUBLIC_SEARCH_PROVIDER_REVALIDATED"
    assert engine.db.events == ["WORKFLOW_PUBLIC_SEARCH_PROVIDER_REVALIDATED"]


def test_research_plan_binds_common_chinese_trace_synonyms():
    from app.skills.research_plan import normalize_and_validate_plan

    plan = {
        "plan_id": "plan-trace-synonym",
        "task_type": "PUBLIC_RESEARCH",
        "research_questions": ["需求条目与源代码之间的追踪关系通常如何恢复？"],
        "queries": ["requirements traceability code retrieval systematic literature"],
        "source_priorities": ["同行评审论文"],
        "time_scope": "2000-2026",
        "evidence_requirements": [],
        "prohibited_inferences": [],
    }
    normalized, validation = normalize_and_validate_plan(plan, strict=True)
    assert validation["status"] == "PASS"
    assert normalized["query_items"][0]["linked_question_indexes"] == [0]


def test_argument_role_vocabulary_is_shared_across_planning_and_writing_schemas():
    import json
    from pathlib import Path

    from app.pack import PromptPack
    from app.proposal_quality import PLAN_ARGUMENT_ROLES

    root = Path(__file__).resolve().parents[1] / "prompt_pack"
    role_schema = json.loads((root / "schemas/common/argument_role.schema.json").read_text(encoding="utf-8"))
    canonical = set(role_schema["enum"])
    required_by_all_fourteen_sections = {
        "PROBLEM", "CENTRAL_CLAIM", "METHOD", "CONTRIBUTION", "CONTEXT", "EVIDENCE",
        "LIMITATION_MECHANISM", "GAP", "RESEARCH_QUESTION", "COMPARISON", "COUNTERARGUMENT",
        "OBJECTIVE", "EVALUATION", "WORK_PACKAGE", "WARRANT", "OUTPUT", "TECHNICAL_DIFFICULTY",
        "BOUNDARY", "INPUT", "DATA_FLOW", "FEEDBACK", "FORMALIZATION", "MECHANISM", "ALGORITHM",
        "DEGRADATION_BASELINE", "DATASET", "BASELINE", "SPLIT", "METRIC", "ABLATION",
        "VALIDITY_THREAT", "CLOSEST_WORK", "NEW_MECHANISM", "COMPARISON_RULE", "MILESTONE",
        "DEPENDENCY", "DELIVERABLE", "RISK_CONTROL", "ACCEPTANCE", "CONFIRMED_CAPABILITY",
        "FEASIBILITY_WARRANT", "LIMITATION", "COMPENSATION_PLAN", "RISK", "MITIGATION", "SYNTHESIS",
        "RQ_CLOSURE",
    }
    assert required_by_all_fourteen_sections <= canonical
    assert canonical == PLAN_ARGUMENT_ROLES

    pack = PromptPack(root)
    prompts = [
        "P-REVISION-PLAN", "P-WRITE-BLUEPRINT", "P-WRITE-BLUEPRINT-CRITIC",
        "P-WRITE-CONTENT", "P-WRITE-CRITIC", "P-EXPRESSION-POLISH", "P-EXPRESSION-CRITIC",
    ]

    def enum_sets(node):
        found = []
        if isinstance(node, dict):
            if node.get("type") == "string" and isinstance(node.get("enum"), list):
                values = set(node["enum"])
                if values & required_by_all_fourteen_sections:
                    found.append(values)
            for value in node.values():
                found.extend(enum_sets(value))
        elif isinstance(node, list):
            for value in node:
                found.extend(enum_sets(value))
        return found

    for prompt_id in prompts:
        for kind in ("input", "output"):
            schema = pack.inlined_schema(prompt_id, kind)
            for values in enum_sets(schema):
                if values & {"INPUT", "MILESTONE", "CONFIRMED_CAPABILITY", "RQ_CLOSURE"}:
                    assert values == canonical


def test_live_context_scope_isolates_concurrent_section_dependencies(monkeypatch):
    import json
    import threading
    from concurrent.futures import ThreadPoolExecutor
    from types import MethodType

    from app.runtime_context import LiveContextBuilder

    candidates = {
        ("wf-a", "sec-a"): {"candidate_id": "candidate-a"},
        ("wf-b", "sec-b"): {"candidate_id": "candidate-b"},
    }

    class FakeDB:
        def fetchall(self, sql, params):
            project_id, workflow_id, prompt_id = params
            assert project_id == "project-1"
            assert prompt_id == "P-WRITE-CONTENT"
            section_id = "sec-a" if workflow_id == "wf-a" else "sec-b"
            return [{
                "input_json": json.dumps({"payload": {"source_section": {"section_id": section_id}}}),
                "output_json": json.dumps({"result": candidates[(workflow_id, section_id)]}),
            }]

        def fetchone(self, sql, params):
            # A missing workflow scope would fall through to the project-global
            # artifact lookup and expose this deliberately wrong candidate.
            return {"content_json": json.dumps({"result": {"candidate_id": "global-wrong"}})}

    builder = object.__new__(LiveContextBuilder)
    builder.runtime_mode = "LIVE"
    builder.db = FakeDB()
    barrier = threading.Barrier(2)

    def fake_build_live(self, prompt_id, project_id, *, workflow_id=None, workflow_state=None, overrides=None):
        barrier.wait(timeout=5)
        return self._result(project_id, "P-WRITE-CONTENT")

    builder._build_live = MethodType(fake_build_live, builder)

    def build_one(workflow_id, section_id):
        return builder.build(
            "P-EXPRESSION-CRITIC",
            "project-1",
            workflow_id=workflow_id,
            workflow_state={"active_section_id": section_id},
        )

    with ThreadPoolExecutor(max_workers=2) as pool:
        results = list(pool.map(lambda pair: build_one(*pair), [("wf-a", "sec-a"), ("wf-b", "sec-b")]))

    assert {item["candidate_id"] for item in results} == {"candidate-a", "candidate-b"}


def test_empty_scoped_project_subgraph_keeps_strict_prompt_shape():
    from app.context_base import ContextBuilder

    project_definition = {
        "items": [{"item_id": "project-item-1", "item_type": "METHOD"}],
        "relations": [],
    }
    contract = {
        "must_advance_claim_ids": ["argument-node-not-a-project-item"],
        "must_use_evidence_ids": ["claim-not-a-project-item"],
    }
    scoped = ContextBuilder._scoped_project_subgraph(project_definition, contract)
    assert scoped == {"item_ids": [], "relation_ids": [], "items": [], "relations": []}
    assert "missing_seed_ids" not in scoped


def test_system_manual_detector_does_not_flag_traceability_reference_titles():
    from app.proposal_quality import _meta_term_hits

    bibliography = (
        "Recovering Traceability Links between Code and Documentation. "
        "Natural Language Processing for Requirements Traceability. "
        "Tracing Requirements and Source Code During Development."
    )
    assert _meta_term_hits(bibliography) == 0
    assert _meta_term_hits("Prompt Trace Gate") == 3


def test_scheme_constraints_preserve_reference_page_exclusion():
    scheme = {
        "rules": [{
            "rule_id": "r-pages",
            "mandatory": True,
            "rule_type": "PAGE_OR_WORD_LIMIT",
            "statement": "正文目标约18页，允许范围16—20页；参考文献页不计入正文页数。",
        }]
    }
    constraints = extract_hard_constraints(scheme)
    assert constraints["main_body_pages"] == {"min": 16, "max": 20}
    assert constraints["references_excluded_from_main_body_pages"] is True
    assert constraints["source_rule_ids"] == ["r-pages"]


def test_main_body_page_metrics_exclude_front_matter_and_reference_only_pages():
    validator = PostExportDeliveryValidator(SimpleNamespace())
    titles = ["摘要", "研究内容", "结论"]
    pages = [
        "项目申请书 封面",
        "目录 摘要 研究内容 结论\n摘要\n" + "正文" * 30,
        "摘要续文 " + "内容" * 30,
        "研究内容 " + "方法" * 30,
        "结论 " + "结论内容" * 20 + "\n参考文献\n[1] source",
        "[2] source\n[3] source",
    ]
    metrics = validator._main_body_page_metrics(
        pages,
        expected_titles=titles,
        references_excluded=True,
    )
    assert metrics == {
        "total_pdf_page_count": 6,
        "main_body_page_count": 4,
        "main_body_start_page": 2,
        "main_body_end_page": 5,
        "reference_start_page": 5,
        "references_excluded": True,
        "measurement_mode": "BODY_BOUNDARIES_FROM_HEADINGS",
    }


def test_main_body_page_metrics_skip_pure_contents_page():
    validator = PostExportDeliveryValidator(SimpleNamespace())
    titles = ["摘要", "背景", "现状", "目标", "内容", "方法"]
    pages = [
        "目录\n" + "\n".join(titles),
        "摘要\n" + "正文" * 40,
        "方法\n" + "正文" * 40,
        "参考文献\n[1] source",
    ]
    metrics = validator._main_body_page_metrics(
        pages,
        expected_titles=titles,
        references_excluded=True,
    )
    assert metrics["main_body_start_page"] == 2
    assert metrics["main_body_end_page"] == 3
    assert metrics["main_body_page_count"] == 2
