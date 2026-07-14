#!/usr/bin/env python3
from __future__ import annotations

import argparse
import asyncio
import csv
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from collections import Counter
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from app.config import Settings
from app.context import ContextBuilder
from app.db import Database
from app.documents import parse_document
from app.executor import PromptExecutor
from app.exporter import DocxExporter
from app.llm import ModelGateway
from app.pack import PromptPack
from app.research import PublicResearchService
from app.security import SecurityRouter
from app.util import new_id, sha256_bytes, utc_now, write_json
from app.workflows import WorkflowEngine
from app.logistics_application_content import SECTION_TITLES, REF_CATALOG

WORKFLOWS = ["WF-1_PROJECT_INTAKE", "WF-2_TEMPLATE_EXTRACTION", "WF-3_HYBRID_ONLINE_ASSIST", "WF-4_PROPOSAL_AUTHORING", "WF-5_SECURITY_REVIEW_AND_EXPORT"]


def build_runtime(output_dir: Path):
    os.environ["MODEL_RUNTIME_MODE"] = "SIMULATED"
    os.environ["APP_DATA_DIR"] = str(output_dir)
    os.environ["PROMPT_PACK_DIR"] = str(ROOT / "prompt_pack")
    settings = Settings.load()
    pack = PromptPack(settings.prompt_pack_dir)
    db = Database(settings.db_path)
    router = SecurityRouter(pack)
    gateway = ModelGateway(settings, pack)
    builder = ContextBuilder(db, pack)
    executor = PromptExecutor(db, pack, router, gateway)
    engine = WorkflowEngine(db, pack, builder, executor, PublicResearchService(settings))
    return settings, pack, db, engine, DocxExporter(db, settings)


def create_project(db: Database, project_doc: dict[str, Any]) -> str:
    project_id = new_id("project")
    now = utc_now()
    db.execute(
        "INSERT INTO projects(id,name,description,security_level,config_json,created_at,updated_at) VALUES(?,?,?,?,?,?,?)",
        (project_id, project_doc["name"], project_doc["description"], project_doc["security_level"], json.dumps(project_doc["config"], ensure_ascii=False), now, now),
    )
    return project_id


def upload_material(settings: Settings, db: Database, project_id: str, path: Path, role: str, security_level: str) -> None:
    raw = path.read_bytes()
    parsed = parse_document(path.name, raw, role, security_level)
    stored = settings.uploads_dir / path.name
    stored.write_bytes(raw)
    db.execute(
        "INSERT INTO documents(id,project_id,filename,role,security_level,document_hash,file_path,parsed_json,created_at) VALUES(?,?,?,?,?,?,?,?,?)",
        (parsed["document_id"], project_id, path.name, role, security_level, parsed["document_hash"], str(stored), json.dumps(parsed, ensure_ascii=False), utc_now()),
    )


def load_materials(settings: Settings, db: Database, project_id: str, materials_dir: Path) -> dict[str, Any]:
    control = materials_dir / "03_control_and_expected"
    expected = json.loads((control / "expected_results.json").read_text(encoding="utf-8"))
    with (control / "upload_manifest.csv").open(encoding="utf-8-sig", newline="") as f:
        rows = list(csv.DictReader(f))
    for row in rows:
        filename = row["filename"]
        candidates = [materials_dir / "01_upload_required" / filename, materials_dir / "02_upload_optional" / filename]
        source = next((p for p in candidates if p.exists()), None)
        if source is None:
            if row["required"].lower() == "true":
                raise FileNotFoundError(filename)
            continue
        upload_material(settings, db, project_id, source, row["role"], row["security_level"])
    return expected


async def finish(engine: WorkflowEngine, project_id: str, workflow_type: str) -> dict[str, Any]:
    workflow = engine.start(project_id, workflow_type)
    for _ in range(1000):
        workflow = await engine.advance(workflow["id"])
        if workflow["status"] == "WAITING_GATE":
            open_gates = [g for g in engine.list_gates(workflow_id=workflow["id"]) if g["status"] == "OPEN"]
            if not open_gates:
                raise RuntimeError(f"{workflow_type}: WAITING_GATE without open gate")
            gate = open_gates[0]
            action = "APPROVE" if "APPROVE" in gate["allowed_actions"] else "CONFIRM"
            engine.decide_gate(gate["id"], action=action, decided_by="complex-e2e", decided_role=gate["required_role"], comment="确定性端到端测试自动批准")
            continue
        if workflow["status"] in {"COMPLETED", "BLOCKED", "CANCELLED"}:
            break
    if workflow["status"] != "COMPLETED":
        raise RuntimeError(f"{workflow_type} failed: {workflow['state'].get('last_error')}")
    return workflow


def dump_rows(path: Path, rows: list[dict[str, Any]]) -> None:
    path.write_text(json.dumps(rows, ensure_ascii=False, indent=2), encoding="utf-8")


def export_trace_bundle(settings: Settings, db: Database, project_id: str, target_dir: Path, final_document: Path) -> Path:
    target_dir.mkdir(parents=True, exist_ok=True)
    shutil.copy2(settings.db_path, target_dir / "proposal_agents.sqlite3")
    tables = {
        "prompt_runs.json": db.fetchall("SELECT * FROM prompt_runs WHERE project_id=? ORDER BY created_at,id", (project_id,)),
        "workflows.json": db.fetchall("SELECT * FROM workflows WHERE project_id=? ORDER BY created_at,id", (project_id,)),
        "gates.json": db.fetchall("SELECT * FROM gates WHERE project_id=? ORDER BY created_at,id", (project_id,)),
        "artifacts.json": db.fetchall("SELECT * FROM artifacts WHERE project_id=? ORDER BY created_at,id", (project_id,)),
        "audit_events.json": db.fetchall("SELECT * FROM audit_events WHERE project_id=? ORDER BY created_at,id", (project_id,)),
        "documents.json": db.fetchall("SELECT * FROM documents WHERE project_id=? ORDER BY created_at,id", (project_id,)),
    }
    for name, rows in tables.items():
        dump_rows(target_dir / name, rows)
    trace_dir = target_dir / "prompt_traces"
    trace_dir.mkdir(exist_ok=True)
    traces = [r for r in tables["artifacts.json"] if r["artifact_type"] == "PROMPT_TRACE"]
    for index, row in enumerate(traces, 1):
        payload = json.loads(row["content_json"])
        (trace_dir / f"{index:04d}_{row['prompt_id']}_v{row['version']}_{row['status']}.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    # Store a concise run index, without duplicating full content in ordinary logs.
    index_rows = [{"index": i, "prompt_id": r["prompt_id"], "status": r["status"], "model_id": r["model_id"], "endpoint_id": r["endpoint_id"], "duration_ms": r["duration_ms"], "input_hash": r["input_hash"], "output_hash": r["output_hash"]} for i, r in enumerate(tables["prompt_runs.json"], 1)]
    dump_rows(target_dir / "prompt_run_index.json", index_rows)
    shutil.copy2(final_document, target_dir / final_document.name)
    zip_path = target_dir.with_suffix(".zip")
    if zip_path.exists(): zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for p in target_dir.rglob("*"):
            if p.is_file(): zf.write(p, arcname=p.relative_to(target_dir).as_posix())
    return zip_path


def convert_pdf(docx_path: Path) -> Path:
    soffice = shutil.which("libreoffice") or shutil.which("soffice")
    if not soffice:
        raise RuntimeError("LibreOffice unavailable; cannot render PDF/page count")
    with tempfile.TemporaryDirectory(prefix="proposal-lo-") as profile_dir:
        profile_uri = Path(profile_dir).resolve().as_uri()
        proc = subprocess.run([soffice, f"-env:UserInstallation={profile_uri}", "--headless", "--convert-to", "pdf", "--outdir", str(docx_path.parent), str(docx_path)], text=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=180)
    pdf = docx_path.with_suffix(".pdf")
    if proc.returncode != 0 or not pdf.exists():
        raise RuntimeError(f"PDF conversion failed: {proc.stdout}\n{proc.stderr}")
    return pdf


def normalized_paragraph(text: str) -> str:
    return re.sub(r"\s+", "", text.strip())


def quality_report(docx_path: Path, pdf_path: Path, db: Database, project_id: str, expected: dict[str, Any]) -> dict[str, Any]:
    doc = Document(str(docx_path))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    headings = [p.text.strip() for p in doc.paragraphs if p.style and (p.style.name.startswith("Heading") or p.style.name.startswith("标题"))]
    refs = re.findall(r"^\[(\d+)\]", all_text, flags=re.M)
    citations = re.findall(r"\[(\d+)\]", all_text)
    substantive = [normalized_paragraph(p.text) for p in doc.paragraphs if len(normalized_paragraph(p.text)) >= 80]
    counts = Counter(substantive)
    duplicates = {k: v for k, v in counts.items() if v > 1}
    pages = len(PdfReader(str(pdf_path)).pages)
    prompt_runs = db.fetchall("SELECT * FROM prompt_runs WHERE project_id=? ORDER BY created_at,id", (project_id,))
    trace_rows = db.fetchall("SELECT * FROM artifacts WHERE project_id=? AND artifact_type='PROMPT_TRACE' ORDER BY created_at,id", (project_id,))
    distinct_prompts = sorted(set(row["prompt_id"] for row in prompt_runs))
    traces_complete = True
    trace_errors = []
    for row in trace_rows:
        payload = json.loads(row["content_json"])
        required = ["prompt_id", "status", "duration_ms", "system_prompt", "input_envelope", "output_schema", "raw_response_text", "environment", "model_id", "endpoint_id"]
        missing = [key for key in required if payload.get(key) is None]
        # Errors before routing/model invocation may legitimately lack some fields.
        if row["status"] != "ERROR" and missing:
            traces_complete = False
            trace_errors.append({"prompt_id": row["prompt_id"], "missing": missing})
    online_inputs = [json.loads(r["input_json"]) for r in prompt_runs if json.loads(r["input_json"]).get("security_context", {}).get("required_environment") == "ONLINE_PUBLIC"]
    online_serialized = json.dumps(online_inputs, ensure_ascii=False)
    privacy_hits = [v for v in expected["privacy_values_not_allowed_online"] if v in online_serialized]
    coverage = {pid: sum(1 for r in prompt_runs if r["prompt_id"] == pid) for pid in expected["required_prompts"]}
    missing_sections = [title for title in SECTION_TITLES if title not in headings]
    missing_fig_captions = [caption for caption in ["业务闭环图", "逻辑结构图", "研究现状与项目切入点图", "目标—内容—技术—成果映射图", "知识图谱模式图", "多智能体协同与门禁关系图", "总体技术路线图", "动态重规划闭环图", "部署架构图", "分层评估与验证框架图", "进度与里程碑图"] if caption not in all_text]
    required_phrases = ["国内外研究现状", "关键技术一：任务语义理解与知识建模", "技术路线", "参考文献", "Prompt、日志与Trace留存规范"]
    report = {
        "pages": pages,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "images": len(doc.inline_shapes),
        "headings": len(headings),
        "text_characters": len(all_text),
        "reference_entries": len(set(refs)),
        "citation_numbers": len(set(citations)),
        "missing_sections": missing_sections,
        "missing_figure_captions": missing_fig_captions,
        "duplicate_substantive_paragraphs": len(duplicates),
        "duplicate_examples": list(duplicates.items())[:5],
        "prompt_runs": len(prompt_runs),
        "distinct_prompts": distinct_prompts,
        "prompt_coverage_counts": coverage,
        "targeted_repair_runs": coverage.get("P-TARGETED-REPAIR", 0),
        "trace_artifacts": len(trace_rows),
        "trace_matches_run_count": len(trace_rows) == len(prompt_runs),
        "trace_payload_complete": traces_complete,
        "trace_errors": trace_errors,
        "online_prompt_runs": len(online_inputs),
        "online_privacy_hits": privacy_hits,
        "required_phrases_present": {phrase: phrase in all_text for phrase in required_phrases},
        "old_sensitive_test_terms_present": any(term in all_text for term in ["单兵", "部队番号", "部署地点", "武器功能"]),
    }
    report["checks"] = {
        "minimum_pages": pages >= expected["minimum_pages"],
        "formal_sections": not missing_sections,
        "minimum_references": report["reference_entries"] >= expected["minimum_references"],
        "minimum_figures": report["images"] >= expected["minimum_figures"],
        "all_prompts_invoked": all(coverage.get(pid, 0) >= 1 for pid in expected["required_prompts"]),
        "targeted_repair_invoked": report["targeted_repair_runs"] >= 1,
        "trace_one_per_run": report["trace_matches_run_count"],
        "trace_complete": report["trace_payload_complete"],
        "privacy_clean": not privacy_hits,
        "no_old_sensitive_terms": not report["old_sensitive_test_terms_present"],
        "no_exact_duplicate_substantive_paragraphs": report["duplicate_substantive_paragraphs"] == 0,
        "figures_complete": not missing_fig_captions,
        "required_content": all(report["required_phrases_present"].values()),
    }
    report["pass"] = all(report["checks"].values())
    return report


async def run(materials_dir: Path, output_dir: Path) -> dict[str, Any]:
    if output_dir.exists(): shutil.rmtree(output_dir)
    output_dir.mkdir(parents=True)
    subprocess.run([sys.executable, str(ROOT / "scripts" / "generate_logistics_figures.py")], check=True)
    settings, pack, db, engine, exporter = build_runtime(output_dir)
    project_doc = json.loads((materials_dir / "03_control_and_expected" / "project_create.json").read_text(encoding="utf-8"))
    project_id = create_project(db, project_doc)
    expected = load_materials(settings, db, project_id, materials_dir)
    workflow_results = {}
    for workflow_type in WORKFLOWS:
        workflow_results[workflow_type] = (await finish(engine, project_id, workflow_type))["status"]
    document = exporter.export(project_id)
    package = exporter.export_package(project_id, document)
    pdf = convert_pdf(document)
    qreport = quality_report(document, pdf, db, project_id, expected)
    trace_bundle = export_trace_bundle(settings, db, project_id, settings.exports_dir / f"{document.stem}_trace_bundle", document)
    report = {
        "schema_version": "1.0",
        "project_id": project_id,
        "runtime_mode": "SIMULATED",
        "workflows": workflow_results,
        "document": str(document),
        "pdf": str(pdf),
        "document_package": str(package),
        "trace_bundle": str(trace_bundle),
        "document_sha256": sha256_bytes(document.read_bytes()),
        "pdf_sha256": sha256_bytes(pdf.read_bytes()),
        "quality": qreport,
        "pass": all(status == "COMPLETED" for status in workflow_results.values()) and qreport["pass"],
    }
    write_json(output_dir / "complex_e2e_report.json", report)
    if not report["pass"]:
        raise RuntimeError(json.dumps(report, ensure_ascii=False, indent=2))
    return report


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--materials-dir", type=Path, default=Path("/mnt/data/logistics_agent_materials_v1"))
    parser.add_argument("--output-dir", type=Path, default=Path("/mnt/data/logistics_agent_final_run"))
    args = parser.parse_args()
    print(json.dumps(asyncio.run(run(args.materials_dir.resolve(), args.output_dir.resolve())), ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
