from __future__ import annotations

import argparse
import json
import shutil
from dataclasses import replace
from pathlib import Path

from docx import Document

from app.config import Settings
from app.delivery_validator import DeliveryValidator
from app.exporter_render import ExportRenderMixin
from app.pdf_exporter import PdfConverter
from app.skills.base import SkillContext
from app.skills.mermaid import MermaidRenderSkill
from app.util import sha256_bytes, utc_now, write_json


DIAGRAMS = [
    (
        "argument-architecture",
        "论证架构图",
        "flowchart LR\n  A[中心命题] --> B[研究问题]\n  B --> C[方法]\n  C --> D[验证]\n  D --> E[贡献]",
    ),
    (
        "technical-route",
        "技术路线图",
        "flowchart TB\n  A[需求分析] --> B[模型构建]\n  B --> C[算法设计]\n  C --> D[原型实现]\n  D --> E[实验验收]",
    ),
    (
        "system-architecture",
        "系统架构图",
        "flowchart LR\n  U[用户材料] --> P[智能体编排]\n  P --> R[研究与图形工具]\n  P --> Q[质量门]\n  Q --> X[DOCX/PDF交付]",
    ),
]


class _Renderer(ExportRenderMixin):
    def __init__(self, settings):
        self.settings = settings


def run(output_dir: Path) -> dict:
    output_dir.mkdir(parents=True, exist_ok=True)
    base = Settings.load()
    settings = replace(
        base,
        data_dir=output_dir,
        db_path=output_dir / "proposal_agents.sqlite3",
        uploads_dir=output_dir / "uploads",
        exports_dir=output_dir / "exports",
        mermaid_browser_executable=(
            base.mermaid_browser_executable
            or shutil.which("chromium")
            or shutil.which("chromium-browser")
            or shutil.which("google-chrome")
            or ""
        ),
    )
    settings.exports_dir.mkdir(parents=True, exist_ok=True)
    skill = MermaidRenderSkill(settings)
    context = SkillContext(
        project_id="d-track-acceptance",
        workflow_id="wf-d-track",
        data_dir=output_dir,
        security_level="INTERNAL",
    )
    markers: list[str] = []
    render_records: list[dict] = []
    try:
        for section_id, caption, source in DIAGRAMS:
            payload = {
                "section_id": section_id,
                "caption": caption,
                "width_cm": 13.5,
                "mermaid_source": source,
            }
            first = skill.run(payload, context)
            second = skill.run(payload, context)
            if not second.output.get("cache_hit"):
                raise RuntimeError(f"Mermaid repeat render did not hit verified cache: {section_id}")
            for key in ("source_sha256", "svg_sha256", "png_sha256"):
                if first.output[key] != second.output[key]:
                    raise RuntimeError(f"Mermaid repeat render hash drift: {section_id}/{key}")
            markers.append(second.output["figure_marker"])
            render_records.append({
                "section_id": section_id,
                "source_sha256": second.output["source_sha256"],
                "svg_sha256": second.output["svg_sha256"],
                "png_sha256": second.output["png_sha256"],
                "cache_hit": second.output["cache_hit"],
            })
    finally:
        skill.close()

    document = Document()
    renderer = _Renderer(settings)
    document.add_heading("D 轨道交付验收", level=1)
    document.add_paragraph("本文件用于验证 Mermaid、图形协议、DOCX、PDF、结构检查和页面视觉检查。")
    for index, ((_, caption, _), marker) in enumerate(zip(DIAGRAMS, markers), 1):
        document.add_heading(f"{index}. {caption}", level=1)
        document.add_paragraph(f"下图给出{caption}的可重复渲染结果。")
        renderer._append_block(document, marker)
    document.add_heading("4. 表格、公式与引用", level=1)
    document.add_paragraph("以下内容验证表格、公式与参考文献在 DOCX/PDF 中保持完整[1]。")
    renderer._append_block(document, "[[TABLE]]指标|验收要求\nMermaid|三类图可重复渲染\n交付物|DOCX/PDF均通过验证")
    renderer._append_block(document, "[[FORMULA]]Q = w_1 S + w_2 C + w_3 V")
    renderer._append_block(document, "[[REFERENCE]][1] D 轨道确定性交付验收基线，2026。")
    renderer._add_page_numbers(document)

    docx_path = settings.exports_dir / "d-track-acceptance.docx"
    document.save(docx_path)
    pdf_path = PdfConverter(settings).convert(docx_path)
    validator = DeliveryValidator(settings)
    validation = validator.validate(
        docx_path,
        pdf_path,
        expected_sections=[
            "D 轨道交付验收",
            "1. 论证架构图",
            "2. 技术路线图",
            "3. 系统架构图",
            "4. 表格、公式与引用",
        ],
        screenshots_dir=output_dir / "page_screenshots",
    )
    validator.require_pass(validation)
    report = {
        "schema_version": "1.0",
        "status": "PASS",
        "created_at": utc_now(),
        "docx": str(docx_path),
        "docx_sha256": sha256_bytes(docx_path.read_bytes()),
        "pdf": str(pdf_path),
        "pdf_sha256": sha256_bytes(pdf_path.read_bytes()),
        "render_records": render_records,
        "delivery_validation": validation,
    }
    report_path = output_dir / "D_TRACK_ACCEPTANCE.json"
    write_json(report_path, report)
    return report


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", type=Path, default=Path("recovery_evidence/d/local"))
    args = parser.parse_args()
    print(json.dumps(run(args.output_dir.resolve()), ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
