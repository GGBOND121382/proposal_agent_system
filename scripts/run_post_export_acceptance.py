from __future__ import annotations

import argparse
import json
import os
import platform
import shutil
import subprocess
import sys
from pathlib import Path
from typing import Any

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from app.config import Settings
from app.db import Database
from app.exporter import DocxExporter
from app.post_export_acceptance import PostExportAcceptanceManager
from app.post_export_validator import PostExportDeliveryValidator
from app.util import new_id, sha256_bytes, sha256_json, utc_now, write_json
from scripts.run_full_proposal_concurrent_acceptance import (
    SECTION_TITLES,
    _export_prompt_evidence,
    _material_manifest,
)


def _source_commit() -> str:
    value = os.getenv("GITHUB_SHA")
    if value:
        return value
    completed = subprocess.run(
        ["git", "rev-parse", "HEAD"], cwd=ROOT, text=True, capture_output=True, check=False
    )
    return completed.stdout.strip() if completed.returncode == 0 else "UNAVAILABLE"


def _approve_export_gates(db: Database, project_id: str, workflow_id: str) -> None:
    now = utc_now()
    for gate_type in ("FINAL_CONTENT_SECURITY_APPROVAL", "FINAL_EXPORT_APPROVAL"):
        db.execute(
            """INSERT INTO gates(
                   id,project_id,workflow_id,gate_type,target_id,target_version,context_hash,
                   question_version,required_role,allowed_actions_json,questions_json,
                   security_level,status,decision_json,created_at,updated_at
               ) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (
                new_id("gate"),
                project_id,
                workflow_id,
                gate_type,
                "post-export-fixed-material",
                1,
                sha256_json({"project_id": project_id, "gate_type": gate_type}),
                1,
                "FIXED_MATERIAL_ACCEPTANCE",
                json.dumps(["APPROVE"], ensure_ascii=False),
                json.dumps(["复用固定材料验收检查点，不修改候选正文。"], ensure_ascii=False),
                "INTERNAL",
                "APPROVED",
                json.dumps({"action": "APPROVE", "scope": "FIXED_MATERIAL_ACCEPTANCE"}, ensure_ascii=False),
                now,
                now,
            ),
        )


def _content_fault_copy(source: Path, destination: Path, *, target_title: str) -> None:
    document = Document(str(source))
    active = False
    for paragraph in document.paragraphs:
        title = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name in {"Heading 1", "标题 1"}:
            active = title == target_title
            continue
        if active and title:
            paragraph.add_run(" TODO")
            document.save(destination)
            return
    raise RuntimeError(f"Could not inject content fault under section: {target_title}")


def _engineering_fault_copy(source: Path, destination: Path, *, missing_title: str) -> None:
    document = Document(str(source))
    for paragraph in document.paragraphs:
        if paragraph.text.strip() != missing_title:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name not in {"Heading 1", "标题 1"}:
            continue
        element = paragraph._element
        element.getparent().remove(element)
        document.save(destination)
        return
    raise RuntimeError(f"Could not remove heading for engineering fault: {missing_title}")


def _owner_set(report: dict[str, Any]) -> set[str]:
    return {str(item.get("owner") or "") for item in report.get("findings") or []}


def _progress(message: str) -> None:
    print(f"[POST_EXPORT] {message}", flush=True)


def _bootstrap_full_proposal(output_dir: Path) -> tuple[Settings, Database, str, dict[str, Any], dict[str, Any], dict[str, Any]]:
    """Create the fixed 14-section checkpoint in a separate process.

    The authoring acceptance is already a tested, recoverable stage. Running it in
    its own process prevents a completed parent workflow from keeping the caller's
    event loop alive through child-task cleanup. The post-export stage then opens
    the persisted SQLite checkpoint and continues without regenerating any chapter.
    """
    bootstrap_dir = output_dir / "authoring_checkpoint"
    if bootstrap_dir.exists():
        shutil.rmtree(bootstrap_dir)
    bootstrap_dir.mkdir(parents=True, exist_ok=True)
    command = [
        sys.executable,
        str(ROOT / "scripts" / "run_full_proposal_concurrent_acceptance.py"),
        "--output-dir",
        str(bootstrap_dir),
    ]
    _progress("build persisted 14-section authoring checkpoint")
    completed = subprocess.run(
        command,
        cwd=ROOT,
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        timeout=900,
        check=False,
        env={**os.environ, "MODEL_RUNTIME_MODE": "SIMULATED"},
    )
    (output_dir / "authoring-checkpoint.log").write_text(completed.stdout or "", encoding="utf-8")
    if completed.returncode != 0:
        raise RuntimeError(
            "Full-proposal checkpoint bootstrap failed: " + (completed.stdout or "")[-4000:]
        )
    authoring_report_path = bootstrap_dir / "FULL_PROPOSAL_CONCURRENT_ACCEPTANCE.json"
    authoring_report = json.loads(authoring_report_path.read_text(encoding="utf-8"))
    if authoring_report.get("status") != "PASS":
        raise RuntimeError(f"Full-proposal checkpoint did not pass: {authoring_report}")

    runtime_dir = bootstrap_dir / "runtime"
    os.environ["MODEL_RUNTIME_MODE"] = "SIMULATED"
    os.environ["APP_DATA_DIR"] = str(runtime_dir)
    os.environ["PROMPT_PACK_DIR"] = str(ROOT / "prompt_pack")
    settings = Settings.load()
    db = Database(settings.db_path)
    project_id = str(authoring_report["project_id"])
    parent_id = str(authoring_report["parent_workflow_id"])
    parent_row = db.fetchone("SELECT * FROM workflows WHERE id=?", (parent_id,))
    if not parent_row:
        raise RuntimeError(f"Persisted parent workflow is missing: {parent_id}")
    parent = dict(parent_row)
    parent["state"] = json.loads(parent.pop("state_json"))
    if parent.get("status") != "COMPLETED":
        raise RuntimeError(f"Persisted full proposal did not complete: {parent.get('status')}")
    workflow_rows = db.fetchall(
        "SELECT workflow_type,status FROM workflows WHERE project_id=?", (project_id,)
    )
    intake = next(
        (dict(row) for row in workflow_rows if row["workflow_type"] == "WF-1_PROJECT_INTAKE"),
        {"status": "MISSING"},
    )
    template = next(
        (dict(row) for row in workflow_rows if row["workflow_type"] == "WF-2_TEMPLATE_EXTRACTION"),
        {"status": "MISSING"},
    )
    return settings, db, project_id, parent, intake, template


def _run(output_dir: Path) -> dict[str, Any]:
    settings, db, project_id, parent, intake, template = _bootstrap_full_proposal(output_dir)
    _approve_export_gates(db, project_id, parent["id"])

    exporter = DocxExporter(db, settings)
    manager = PostExportAcceptanceManager(db, settings, exporter)
    _progress("run production DOCX/PDF acceptance")
    accepted = manager.run(project_id, workflow_id=parent["id"], reuse_verified=False)
    _progress("production acceptance completed")
    if accepted["status"] != "PASS":
        raise RuntimeError(f"Post-export acceptance failed: {accepted['status']}")

    document_path = Path(accepted["document"]["path"])
    package_path = Path(accepted["package"]["path"])
    candidate_snapshot = exporter.candidate_snapshot(project_id)
    candidates = exporter._candidate_runs(project_id)
    validator = PostExportDeliveryValidator(settings)

    fault_dir = output_dir / "fault_injection"
    fault_dir.mkdir(parents=True, exist_ok=True)

    _progress("run content fault routing check")
    content_docx = fault_dir / "content-fault.docx"
    _content_fault_copy(document_path, content_docx, target_title="技术路线")
    content_pdf = exporter.pdf_converter.convert(content_docx)
    content_report = validator.validate(
        content_docx,
        content_pdf,
        expected_sections=SECTION_TITLES,
        expected_candidates=candidates,
        screenshots_dir=fault_dir / "content-fault-pages",
    )
    write_json(fault_dir / "content-fault-report.json", content_report)

    _progress("run engineering fault routing check")
    engineering_docx = fault_dir / "engineering-fault.docx"
    _engineering_fault_copy(document_path, engineering_docx, missing_title="研究目标")
    engineering_pdf = exporter.pdf_converter.convert(engineering_docx)
    engineering_report = validator.validate(
        engineering_docx,
        engineering_pdf,
        expected_sections=SECTION_TITLES,
        expected_candidates=candidates,
        screenshots_dir=fault_dir / "engineering-fault-pages",
    )
    write_json(fault_dir / "engineering-fault-report.json", engineering_report)

    _progress("run restart reuse check")
    restarted_manager = PostExportAcceptanceManager(db, settings)
    restarted = restarted_manager.run(project_id, workflow_id=parent["id"])

    delivery = json.loads(Path(accepted["delivery_report"]["path"]).read_text(encoding="utf-8"))
    structure = json.loads(Path(accepted["structure_report"]["path"]).read_text(encoding="utf-8"))
    visual = json.loads(Path(accepted["visual_report"]["path"]).read_text(encoding="utf-8"))
    review_history = parent["state"].get("full_proposal_review_history") or []
    final_review = review_history[-1] if review_history else {}

    checks = {
        "prerequisites_completed": intake["status"] == template["status"] == "COMPLETED",
        "full_proposal_completed": parent["status"] == "COMPLETED",
        "full_integration_passed": final_review.get("status") == "PASS",
        "fourteen_reviewed_sections": candidate_snapshot["section_count"] == len(SECTION_TITLES) == 14,
        "export_matches_latest_integration_review": accepted["integration_review"]["status"] == "PASS",
        "docx_pdf_delivery_passed": delivery.get("status") == "PASS",
        "no_structure_findings": structure.get("status") == "PASS" and not structure.get("findings"),
        "no_visual_findings": visual.get("status") == "PASS" and not visual.get("findings"),
        "candidate_parity_complete": all(
            int((structure.get("candidate_parity") or {}).get(key) or 0) == 0
            for key in ("docx_missing_unit_count", "pdf_missing_unit_count")
        ),
        "page_screenshots_complete": len(accepted.get("screenshots") or []) == int(visual.get("page_count") or 0) > 0,
        "content_fault_routes_to_writer": content_report.get("status") == "FAIL"
        and "WRITING_AGENT" in _owner_set(content_report)
        and any(
            item.get("responsible_section_ids")
            for item in content_report.get("findings") or []
            if item.get("owner") == "WRITING_AGENT"
        ),
        "engineering_fault_routes_to_export_engineering": engineering_report.get("status") == "FAIL"
        and "EXPORT_ENGINEERING" in _owner_set(engineering_report)
        and "WRITING_AGENT" not in {
            str(item.get("owner") or "")
            for item in engineering_report.get("findings") or []
            if item.get("code") in {"D5_MISSING_SECTION", "D5_SECTION_ORDER_OR_SET_DRIFT"}
        },
        "production_package_unchanged_by_fault_injection": sha256_bytes(package_path.read_bytes()) == accepted["package"]["sha256"],
        "restart_reused_verified_attempt": restarted.get("reused_after_restart") is True
        and restarted.get("attempt_id") == accepted.get("attempt_id"),
        "quality_ledger_has_no_open_blockers": manager.quality_manager.open_blockers(project_id) == [],
    }

    _progress("write evidence manifests")
    prompt_run_count = _export_prompt_evidence(db, project_id, output_dir)
    material_manifest = _material_manifest(db, project_id)
    report = {
        "schema_version": "1.0",
        "stage": "DOCX_PDF_POST_EXPORT_ACCEPTANCE",
        "status": "PASS" if all(checks.values()) else "FAIL",
        "source_commit": _source_commit(),
        "created_at": utc_now(),
        "project_id": project_id,
        "workflow_id": parent["id"],
        "checks": checks,
        "metrics": {
            "section_count": candidate_snapshot["section_count"],
            "candidate_set_hash": candidate_snapshot["candidate_set_hash"],
            "page_count": visual.get("page_count"),
            "screenshot_count": len(accepted.get("screenshots") or []),
            "prompt_run_count": prompt_run_count,
            "content_fault_finding_count": content_report.get("finding_count"),
            "engineering_fault_finding_count": engineering_report.get("finding_count"),
        },
        "artifacts": {
            "docx": accepted["document"],
            "pdf": accepted["pdf"],
            "package": accepted["package"],
            "post_export_report": accepted["report_path"],
            "delivery_report": accepted["delivery_report"],
            "structure_report": accepted["structure_report"],
            "visual_report": accepted["visual_report"],
            "content_fault_report": str(fault_dir / "content-fault-report.json"),
            "engineering_fault_report": str(fault_dir / "engineering-fault-report.json"),
            "workflow_checkpoint": str(settings.db_path),
        },
        "material_manifest": material_manifest,
        "environment": {
            "python": sys.version,
            "platform": platform.platform(),
            "runtime_mode": settings.runtime_mode,
            "libreoffice": shutil.which("libreoffice") or shutil.which("soffice"),
            "pdftoppm": shutil.which("pdftoppm"),
        },
        "evidence_scope": (
            "固定材料用于验证完整候选溯源、真实 DOCX/PDF 转换、结构与页面验收、"
            "责任分类和重启复用；不作为 G3 LIVE 模型语义能力证明。"
        ),
    }
    write_json(output_dir / "POST_EXPORT_ACCEPTANCE.json", report)
    (output_dir / "acceptance_report.md").write_text(
        "# DOCX/PDF 导出后验收\n\n"
        f"- 状态：**{report['status']}**\n"
        f"- 章节：{report['metrics']['section_count']}\n"
        f"- PDF 页数：{report['metrics']['page_count']}\n"
        f"- 页面截图：{report['metrics']['screenshot_count']}\n"
        f"- Prompt Run：{prompt_run_count}\n"
        f"- 候选集合哈希：`{candidate_snapshot['candidate_set_hash']}`\n"
        "- 正文故障：返回责任 Writing Agent。\n"
        "- 导出故障：返回 Export Engineering，禁止改正文掩盖。\n"
        "- 重启：复用已验证且哈希一致的交付物。\n",
        encoding="utf-8",
    )
    return report


def main() -> int:
    parser = argparse.ArgumentParser(description="Run fixed-material DOCX/PDF post-export acceptance.")
    parser.add_argument("--output-dir", type=Path, required=True)
    args = parser.parse_args()
    output_dir = args.output_dir.resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    report = _run(output_dir)
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0 if report["status"] == "PASS" else 1


if __name__ == "__main__":
    raise SystemExit(main())
