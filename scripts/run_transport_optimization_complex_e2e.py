#!/usr/bin/env python3
from __future__ import annotations

import argparse
import asyncio
import csv
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from collections import Counter
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from app.config import Settings
from app.context import ContextBuilder
from app.db import Database
from app.diagram_enrichment import DiagramEnrichmentService
from app.documents import parse_document
from app.executor import PromptExecutor
from app.exporter import DocxExporter
from app.llm import ModelGateway
from app.pack import PromptPack
from app.research import PublicResearchService
from app.security import SecurityRouter
from app.skill_setup import build_skill_executor
from app.util import new_id, sha256_bytes, utc_now, write_json
from app.workflows import WorkflowEngine
from app.transport_optimization_application_content import SECTION_TITLES

WORKFLOWS = ["WF-1_PROJECT_INTAKE","WF-2_TEMPLATE_EXTRACTION","WF-3_HYBRID_ONLINE_ASSIST","WF-4_PROPOSAL_AUTHORING","WF-5_SECURITY_REVIEW_AND_EXPORT"]


def build_runtime(output_dir: Path, research_file: Path):
    os.environ.update({
        "MODEL_RUNTIME_MODE":"SIMULATED", "APP_DATA_DIR":str(output_dir), "PROMPT_PACK_DIR":str(ROOT/"prompt_pack"),
        "PUBLIC_SEARCH_PROVIDER":"connector", "PUBLIC_RESEARCH_CONNECTOR_FILE":str(research_file), "PUBLIC_SEARCH_MAX_RESULTS":"40",
        "MERMAID_JS_PATH":str(ROOT/"third_party"/"mermaid"/"mermaid.min.js"), "MERMAID_BROWSER_EXECUTABLE":os.getenv("MERMAID_BROWSER_EXECUTABLE","/usr/bin/chromium"),
        "SKILL_TIMEOUT_SECONDS":"60",
    })
    settings=Settings.load(); pack=PromptPack(settings.prompt_pack_dir); db=Database(settings.db_path)
    router=SecurityRouter(pack); gateway=ModelGateway(settings,pack); context=ContextBuilder(db,pack)
    executor=PromptExecutor(db,pack,router,gateway); skills=build_skill_executor(db,settings)
    research=PublicResearchService(settings,skills); diagrams=DiagramEnrichmentService(db,pack,skills)
    engine=WorkflowEngine(db,pack,context,executor,research,diagrams)
    return settings,pack,db,engine,DocxExporter(db,settings)


def create_project(db: Database, doc: dict[str,Any]) -> str:
    pid=new_id("project"); now=utc_now()
    db.execute("INSERT INTO projects(id,name,description,security_level,config_json,created_at,updated_at) VALUES(?,?,?,?,?,?,?)",(pid,doc["name"],doc["description"],doc["security_level"],json.dumps(doc["config"],ensure_ascii=False),now,now))
    return pid


def upload(settings: Settings, db: Database, project_id: str, path: Path, role: str, level: str) -> None:
    raw=path.read_bytes(); parsed=parse_document(path.name,raw,role,level); stored=settings.uploads_dir/path.name; stored.write_bytes(raw)
    db.execute("INSERT INTO documents(id,project_id,filename,role,security_level,document_hash,file_path,parsed_json,created_at) VALUES(?,?,?,?,?,?,?,?,?)",(parsed["document_id"],project_id,path.name,role,level,parsed["document_hash"],str(stored),json.dumps(parsed,ensure_ascii=False),utc_now()))


def load_materials(settings: Settings, db: Database, pid: str, root: Path) -> dict[str,Any]:
    control=root/"03_control_and_expected"
    expected=json.loads((control/"expected_results.json").read_text(encoding="utf-8"))
    with (control/"upload_manifest.csv").open(encoding="utf-8-sig",newline="") as f: rows=list(csv.DictReader(f))
    for row in rows:
        path=next((root/d/row["filename"] for d in ["01_upload_required","02_upload_optional"] if (root/d/row["filename"]).exists()),None)
        if path is None:
            if row["required"].lower()=="true": raise FileNotFoundError(row["filename"])
            continue
        upload(settings,db,pid,path,row["role"],row["security_level"])
    return expected


async def finish(engine: WorkflowEngine, pid: str, kind: str) -> dict[str,Any]:
    wf=engine.start(pid,kind)
    for _ in range(2000):
        wf=await engine.advance(wf["id"])
        if wf["status"]=="WAITING_GATE":
            gates=[g for g in engine.list_gates(workflow_id=wf["id"]) if g["status"]=="OPEN"]
            if not gates: raise RuntimeError(f"{kind}: no open gate")
            gate=gates[0]; action="APPROVE" if "APPROVE" in gate["allowed_actions"] else "CONFIRM"
            engine.decide_gate(gate["id"],action=action,decided_by="transport-e2e",decided_role=gate["required_role"],comment="物流运输优化复杂端到端验证自动批准")
            continue
        if wf["status"] in {"COMPLETED","BLOCKED","CANCELLED"}: break
    if wf["status"]!="COMPLETED": raise RuntimeError(f"{kind} failed: {wf['state'].get('last_error')}")
    return wf


def convert_pdf(docx: Path) -> Path:
    soffice=shutil.which("libreoffice") or shutil.which("soffice")
    if not soffice: raise RuntimeError("LibreOffice unavailable")
    with tempfile.TemporaryDirectory(prefix="transport-proposal-lo-") as tmp:
        proc=subprocess.run([soffice,f"-env:UserInstallation={Path(tmp).resolve().as_uri()}","--headless","--convert-to","pdf","--outdir",str(docx.parent),str(docx)],text=True,stdout=subprocess.PIPE,stderr=subprocess.PIPE,timeout=240)
    pdf=docx.with_suffix(".pdf")
    if proc.returncode or not pdf.exists(): raise RuntimeError(proc.stdout+"\n"+proc.stderr)
    return pdf


def dump(path: Path, data: Any) -> None: path.write_text(json.dumps(data,ensure_ascii=False,indent=2),encoding="utf-8")


def export_audit(settings: Settings, db: Database, pid: str, target: Path, docx: Path, pdf: Path) -> Path:
    target.mkdir(parents=True,exist_ok=True); shutil.copy2(settings.db_path,target/"proposal_agents.sqlite3")
    queries={
        "prompt_runs.json":"SELECT * FROM prompt_runs WHERE project_id=? ORDER BY created_at,id",
        "skill_runs.json":"SELECT * FROM skill_runs WHERE project_id=? ORDER BY created_at,id",
        "workflows.json":"SELECT * FROM workflows WHERE project_id=? ORDER BY created_at,id",
        "gates.json":"SELECT * FROM gates WHERE project_id=? ORDER BY created_at,id",
        "artifacts.json":"SELECT * FROM artifacts WHERE project_id=? ORDER BY created_at,id",
        "audit_events.json":"SELECT * FROM audit_events WHERE project_id=? ORDER BY created_at,id",
        "documents.json":"SELECT * FROM documents WHERE project_id=? ORDER BY created_at,id",
    }
    tables={name:db.fetchall(sql,(pid,)) for name,sql in queries.items()}
    for name,rows in tables.items(): dump(target/name,rows)
    traces=target/"prompt_traces"; traces.mkdir()
    for i,row in enumerate([r for r in tables["artifacts.json"] if r["artifact_type"]=="PROMPT_TRACE"],1):
        dump(traces/f"{i:04d}_{row['prompt_id']}_{row['status']}.json",json.loads(row["content_json"]))
    # Copy complete skill artifact trees so research snapshots and Mermaid source are independently inspectable.
    for dirname in ["research_archive","diagram_artifacts"]:
        src=settings.data_dir/dirname
        if src.exists(): shutil.copytree(src,target/dirname,dirs_exist_ok=True)
    shutil.copy2(docx,target/docx.name); shutil.copy2(pdf,target/pdf.name)
    zip_path=target.with_suffix(".zip")
    if zip_path.exists(): zip_path.unlink()
    with zipfile.ZipFile(zip_path,"w",compression=zipfile.ZIP_DEFLATED) as zf:
        for p in target.rglob("*"):
            if p.is_file(): zf.write(p,arcname=p.relative_to(target).as_posix())
    return zip_path


def quality(docx: Path,pdf: Path,db: Database,pid: str,settings: Settings,expected: dict[str,Any]) -> dict[str,Any]:
    doc=Document(str(docx)); all_text="\n".join(p.text for p in doc.paragraphs)
    headings=[p.text.strip() for p in doc.paragraphs if p.style and (p.style.name.startswith("Heading") or p.style.name.startswith("标题"))]
    refs=set(re.findall(r"^\[(\d+)\]",all_text,flags=re.M)); citations=set(re.findall(r"\[(\d+)\]",all_text))
    normalized=[re.sub(r"\s+","",p.text) for p in doc.paragraphs if len(re.sub(r"\s+","",p.text))>=100]
    dup={k:v for k,v in Counter(normalized).items() if v>1}
    runs=db.fetchall("SELECT * FROM prompt_runs WHERE project_id=? ORDER BY created_at,id",(pid,)); skills=db.fetchall("SELECT * FROM skill_runs WHERE project_id=? ORDER BY created_at,id",(pid,))
    artifacts=db.fetchall("SELECT * FROM artifacts WHERE project_id=? ORDER BY created_at,id",(pid,)); traces=[r for r in artifacts if r["artifact_type"]=="PROMPT_TRACE"]
    coverage={pid_:sum(1 for r in runs if r["prompt_id"]==pid_) for pid_ in expected["required_prompts"]}
    online=json.dumps([json.loads(r["input_json"]) for r in runs if json.loads(r["input_json"]).get("security_context",{}).get("required_environment")=="ONLINE_PUBLIC"],ensure_ascii=False)
    privacy=[x for x in expected["privacy_values_not_allowed_online"] if x in online]
    mmd=list((settings.data_dir/"diagram_artifacts").rglob("*.mmd")); png=list((settings.data_dir/"diagram_artifacts").rglob("*.png")); svg=list((settings.data_dir/"diagram_artifacts").rglob("*.svg"))
    manifests=list((settings.data_dir/"research_archive").rglob("manifest.json")); research_sources=0; research_modes=[]; research_integrity=True
    for m in manifests:
        payload=json.loads(m.read_text(encoding="utf-8")); research_sources+=payload.get("source_count",0); research_modes.append(payload.get("retrieval_mode"))
        for record in payload.get("records",[]):
            for key in ["source_id","url","raw_path","text_path","metadata_path","snapshot_sha256","text_sha256","retrieved_at"]:
                if not record.get(key): research_integrity=False
            for path_key in ["raw_path","text_path","metadata_path"]:
                if record.get(path_key) and not Path(record[path_key]).exists(): research_integrity=False
    trace_complete=True
    for r in traces:
        payload=json.loads(r["content_json"])
        for key in ["prompt_id","system_prompt","input_envelope","output_schema","raw_response_text","environment","model_id","endpoint_id","duration_ms","status"]:
            if r["status"]!="ERROR" and payload.get(key) is None: trace_complete=False
    pages=len(PdfReader(str(pdf)).pages)
    missing_sections=[t for t in SECTION_TITLES if t not in headings]
    result={
        "pages":pages,"text_characters":len(all_text),"paragraphs":len(doc.paragraphs),"tables":len(doc.tables),"images":len(doc.inline_shapes),
        "reference_entries":len(refs),"citation_numbers":len(citations),"missing_sections":missing_sections,"duplicate_substantive_paragraphs":len(dup),
        "prompt_runs":len(runs),"prompt_coverage":coverage,"trace_artifacts":len(traces),"skill_runs":len(skills),
        "skill_counts":{sid:sum(1 for r in skills if r["skill_id"]==sid) for sid in sorted({r["skill_id"] for r in skills})},
        "mermaid_sources":len(mmd),"mermaid_png":len(png),"mermaid_svg":len(svg),"research_sources":research_sources,"research_modes":research_modes,
        "research_archive_integrity":research_integrity,"online_privacy_hits":privacy,
        "required_phrases":{p:p in all_text for p in expected["required_phrases"]},
    }
    result["checks"]={
        "minimum_pages":pages>=expected["minimum_pages"],"all_sections":not missing_sections,"minimum_references":len(refs)>=expected["minimum_references"],
        "minimum_figures":len(doc.inline_shapes)>=expected["minimum_figures"],"minimum_mermaid_sources":len(mmd)>=expected["minimum_mermaid_sources"],
        "diagram_triplets":len(mmd)==len(png)==len(svg),"minimum_research_sources":research_sources>=expected["minimum_research_sources"],
        "real_research_archive":bool(manifests) and all(mode in {"LIVE_CONNECTOR_ARCHIVE","LIVE_SEARXNG"} for mode in research_modes),
        "research_integrity":research_integrity,"all_prompts":all(v>=1 for v in coverage.values()),"targeted_repair":coverage.get("P-TARGETED-REPAIR",0)>=1,
        "trace_one_per_run":len(traces)==len(runs),"trace_complete":trace_complete,"privacy_clean":not privacy,
        "no_exact_duplicate_substantive_paragraphs":not dup,"required_content":all(result["required_phrases"].values()),
        "research_skill_called":result["skill_counts"].get("public_research.archive",0)>=1,"mermaid_skill_called":result["skill_counts"].get("mermaid.render",0)>=expected["minimum_mermaid_sources"],
    }
    result["pass"]=all(result["checks"].values())
    return result


async def run(materials: Path,output: Path) -> dict[str,Any]:
    if output.exists(): shutil.rmtree(output)
    output.mkdir(parents=True)
    control=materials/"03_control_and_expected"; research_file=control/"transport_optimization_connector_response.json"
    settings,pack,db,engine,exporter=build_runtime(output,research_file)
    project_doc=json.loads((control/"project_create.json").read_text(encoding="utf-8")); pid=create_project(db,project_doc)
    expected=load_materials(settings,db,pid,materials)
    workflows={}
    for kind in WORKFLOWS: workflows[kind]=(await finish(engine,pid,kind))["status"]
    docx=exporter.export(pid); package=exporter.export_package(pid,docx); pdf=convert_pdf(docx)
    q=quality(docx,pdf,db,pid,settings,expected)
    audit=export_audit(settings,db,pid,settings.exports_dir/f"{docx.stem}_full_audit",docx,pdf)
    report={"schema_version":"1.0","project_id":pid,"runtime_mode":"SIMULATED","public_research_provider":"connector","workflows":workflows,"document":str(docx),"pdf":str(pdf),"document_package":str(package),"audit_package":str(audit),"document_sha256":sha256_bytes(docx.read_bytes()),"pdf_sha256":sha256_bytes(pdf.read_bytes()),"quality":q}
    report["pass"]=all(v=="COMPLETED" for v in workflows.values()) and q["pass"]
    write_json(output/"transport_optimization_e2e_report.json",report)
    if not report["pass"]: raise RuntimeError(json.dumps(report,ensure_ascii=False,indent=2))
    return report


def main() -> None:
    p=argparse.ArgumentParser(); p.add_argument("--materials-dir",type=Path,default=Path("/mnt/data/transport_optimization_materials_v1")); p.add_argument("--output-dir",type=Path,default=Path("/mnt/data/transport_optimization_final_run")); a=p.parse_args()
    print(json.dumps(asyncio.run(run(a.materials_dir.resolve(),a.output_dir.resolve())),ensure_ascii=False,indent=2))

if __name__=="__main__": main()
