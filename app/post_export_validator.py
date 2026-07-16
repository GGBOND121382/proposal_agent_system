from __future__ import annotations

import re
import shutil
import subprocess
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader

from .delivery_validator_runtime import DeliveryValidator as RuntimeDeliveryValidator
from .figure_protocol import FigureProtocolError, parse_figure_block
from .util import sha256_json, sha256_text, utc_now, write_json


class PostExportDeliveryValidator(RuntimeDeliveryValidator):
    """Delivery validation with candidate provenance and DOCX/PDF parity checks.

    The exporter is allowed to transform layout directives, but it may not silently
    drop reviewed content or invent a different section set. These checks compare
    the final files with the exact Expression-Critic-approved candidate snapshot.
    """

    def validate(
        self,
        docx_path: Path,
        pdf_path: Path,
        *,
        expected_sections: list[str] | None = None,
        expected_candidates: list[dict[str, Any]] | None = None,
        screenshots_dir: Path | None = None,
    ) -> dict[str, Any]:
        docx_path = docx_path.resolve()
        pdf_path = pdf_path.resolve()
        screenshots_dir = screenshots_dir or docx_path.parent / f"{docx_path.stem}-pages"
        structure = self.validate_structure(
            docx_path,
            pdf_path,
            expected_sections=expected_sections or [],
            expected_candidates=expected_candidates or [],
        )
        visual = self.validate_visual(docx_path, pdf_path, screenshots_dir=screenshots_dir)
        findings = structure["findings"] + visual["findings"]
        for finding in findings:
            self._enrich_finding(finding)
        blocking = [item for item in findings if item.get("blocking", True)]
        report = {
            "schema_version": "2.0",
            "validated_at": utc_now(),
            "status": "FAIL" if blocking else "PASS",
            "docx_filename": docx_path.name,
            "docx_sha256": self._sha_file(docx_path),
            "pdf_filename": pdf_path.name,
            "pdf_sha256": self._sha_file(pdf_path),
            "structure_report": structure["report_path"],
            "visual_report": visual["report_path"],
            "screenshot_dir": str(screenshots_dir),
            "finding_count": len(findings),
            "blocking_finding_count": len(blocking),
            "finding_owner_counts": self._owner_counts(findings),
            "findings": findings,
        }
        overall_path = docx_path.with_suffix(".delivery-validation.json")
        write_json(overall_path, report)
        report["report_path"] = str(overall_path)
        return report

    def validate_structure(
        self,
        docx_path: Path,
        pdf_path: Path,
        *,
        expected_sections: list[str],
        expected_candidates: list[dict[str, Any]] | None = None,
    ) -> dict[str, Any]:
        base = super().validate_structure(
            docx_path, pdf_path, expected_sections=expected_sections
        )
        content_codes_replaced = {
            "D5_PLACEHOLDER_BRACES",
            "D5_PLACEHOLDER_WORD",
            "D5_INTERNAL_RUNTIME_TERM",
        }
        findings = [
            item for item in (base.get("findings") or [])
            if str(item.get("code") or "") not in content_codes_replaced
        ]
        document = Document(str(docx_path))
        candidates = expected_candidates or []
        expected_titles = [
            str(item.get("section_title") or "").strip()
            for item in candidates
            if str(item.get("section_title") or "").strip()
        ] or list(expected_sections)
        actual_titles = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
            and self._heading_level(paragraph.style.name if paragraph.style else "") == 1
            and paragraph.text.strip() != "目录"
        ]
        if actual_titles != expected_titles:
            findings.append(
                self._finding(
                    "D5_SECTION_ORDER_OR_SET_DRIFT",
                    "P0",
                    "DOCX 一级章节集合或顺序与最终审查通过的候选集合不一致。",
                    location="DOCX headings",
                    blocking=True,
                    category="FORMAT",
                    target_type="DOCX_STRUCTURE",
                    evidence={"expected": expected_titles, "actual": actual_titles},
                )
            )

        expected = self._expected_visible_manifest(candidates)
        actual_text = self._docx_visible_text(document)
        actual_normalized = self._normalize(actual_text)
        missing_units: list[dict[str, str]] = []
        for unit in expected["units"]:
            # Formula directives are rendered as native OMML equations. python-docx
            # intentionally excludes OMML tokens from paragraph.text, so formula
            # preservation is checked structurally by formula_count below.
            if unit.get("kind") == "FORMULA":
                continue
            normalized = self._normalize(unit["text"])
            if len(normalized) >= 4 and normalized not in actual_normalized:
                missing_units.append(unit)
        if missing_units:
            affected = sorted({item["section_id"] for item in missing_units if item["section_id"]})
            findings.append(
                self._finding(
                    "D5_REVIEWED_CONTENT_LOSS",
                    "P0",
                    f"DOCX 丢失 {len(missing_units)} 个已通过 Expression Critic 的可见内容单元。",
                    location="DOCX body",
                    blocking=True,
                    category="FORMAT",
                    target_type="DOCX_RENDER",
                    responsible_section_ids=affected,
                    evidence={"missing_units": missing_units[:20]},
                )
            )

        actual_formulas = len(document.element.xpath(".//m:oMathPara"))
        if actual_formulas != expected["formula_count"]:
            findings.append(
                self._finding(
                    "D5_FORMULA_COUNT_MISMATCH",
                    "P0",
                    f"公式数量与已审查指令不一致：期望 {expected['formula_count']}，实际 {actual_formulas}。",
                    location="DOCX formulas",
                    blocking=True,
                    category="FORMAT",
                    target_type="DOCX_RENDER",
                    evidence={"expected": expected["formula_count"], "actual": actual_formulas},
                )
            )

        actual_figures = len(document.inline_shapes)
        if actual_figures != expected["figure_count"]:
            findings.append(
                self._finding(
                    "D5_FIGURE_COUNT_MISMATCH",
                    "P0",
                    f"图形数量与已审查指令不一致：期望 {expected['figure_count']}，实际 {actual_figures}。",
                    location="DOCX figures",
                    blocking=True,
                    category="FORMAT",
                    target_type="DOCX_RENDER",
                    evidence={"expected": expected["figure_count"], "actual": actual_figures},
                )
            )
        actual_tables = len(document.tables)
        if actual_tables != expected["table_count"]:
            findings.append(
                self._finding(
                    "D5_TABLE_COUNT_MISMATCH",
                    "P1",
                    f"表格数量与已审查指令不一致：期望 {expected['table_count']}，实际 {actual_tables}。",
                    location="DOCX tables",
                    blocking=True,
                    category="FORMAT",
                    target_type="DOCX_RENDER",
                    evidence={"expected": expected["table_count"], "actual": actual_tables},
                )
            )

        pdf_text = self._pdf_text(pdf_path)
        pdf_normalized = self._normalize(pdf_text)
        # PDF text extractors linearize tables column-by-column or interleave cells,
        # so an individual DOCX cell is not guaranteed to survive as one contiguous
        # text unit even when the rendered table is complete. Table preservation is
        # verified independently by DOCX table structure/count and page visual QA.
        # Keep PDF text-layer parity for prose-like units only.
        pdf_parity_units = [
            unit for unit in expected["units"]
            if unit.get("kind") not in {"TABLE_CELL", "FORMULA"}
        ]
        pdf_missing = [
            unit
            for unit in pdf_parity_units
            if len(self._normalize(unit["text"])) >= 4
            and not self._fuzzy_unit_present(self._normalize(unit["text"]), pdf_normalized)
        ]
        if pdf_missing:
            findings.append(
                self._finding(
                    "D5_PDF_CONTENT_LOSS",
                    "P0",
                    f"PDF 文本层缺少 {len(pdf_missing)} 个 DOCX 中应存在的已审查内容单元。",
                    location="PDF text layer",
                    blocking=True,
                    category="FORMAT",
                    target_type="PDF_RENDER",
                    responsible_section_ids=sorted(
                        {item["section_id"] for item in pdf_missing if item["section_id"]}
                    ),
                    evidence={"missing_units": pdf_missing[:20]},
                )
            )

        section_by_paragraph = self._paragraph_sections(document)
        for code, pattern in self.PLACEHOLDER_PATTERNS.items():
            affected = sorted(
                {
                    section_by_paragraph.get(index, "")
                    for index, paragraph in enumerate(document.paragraphs)
                    if pattern.search(paragraph.text or "")
                    and section_by_paragraph.get(index, "")
                }
            )
            if affected:
                findings.append(
                    self._finding(
                        f"D5_CONTENT_{code}",
                        "P1",
                        "最终正文仍包含占位内容，必须返回责任章节修复后重新执行全文审查。",
                        location="DOCX section body",
                        blocking=True,
                        category="CONTENT",
                        target_type="SECTION_CANDIDATE",
                        responsible_section_ids=affected,
                    )
                )
        runtime_pattern = self.INTERNAL_PATTERNS["INTERNAL_RUNTIME_TERM"]
        affected_runtime = sorted(
            {
                section_by_paragraph.get(index, "")
                for index, paragraph in enumerate(document.paragraphs)
                if runtime_pattern.search(paragraph.text or "")
                and section_by_paragraph.get(index, "")
            }
        )
        if affected_runtime:
            findings.append(
                self._finding(
                    "D5_CONTENT_DOCUMENT_TYPE_DRIFT",
                    "P1",
                    "主文出现运行时或测试术语，属于正文文种漂移。",
                    location="DOCX section body",
                    blocking=True,
                    category="CONTENT",
                    target_type="SECTION_CANDIDATE",
                    responsible_section_ids=affected_runtime,
                )
            )

        for finding in findings:
            self._enrich_finding(finding)
        structure = dict(base)
        structure.update(
            {
                "schema_version": "2.0",
                "status": "FAIL" if any(item.get("blocking", True) for item in findings) else "PASS",
                "findings": findings,
                "candidate_parity": {
                    "candidate_set_hash": sha256_json(
                        {
                            "sections": [
                                {
                                    "section_id": item.get("section_id"),
                                    "candidate_id": item.get("candidate_id"),
                                    "paragraphs": item.get("paragraphs") or [],
                                }
                                for item in candidates
                            ]
                        }
                    ),
                    "expected_section_count": len(expected_titles),
                    "actual_section_count": len(actual_titles),
                    "expected_visible_unit_count": len(expected["units"]),
                    "docx_missing_unit_count": len(missing_units),
                    "pdf_missing_unit_count": len(pdf_missing),
                    "expected_figure_count": expected["figure_count"],
                    "actual_figure_count": actual_figures,
                    "expected_table_count": expected["table_count"],
                    "actual_table_count": actual_tables,
                    "source_visible_sha256": sha256_text(
                        "\n".join(item["text"] for item in expected["units"])
                    ),
                    "docx_visible_sha256": sha256_text(actual_text),
                    "pdf_text_sha256": sha256_text(pdf_text),
                },
            }
        )
        path = docx_path.with_suffix(".structure-findings.json")
        write_json(path, structure)
        structure["report_path"] = str(path)
        return structure

    def _expected_visible_manifest(self, candidates: list[dict[str, Any]]) -> dict[str, Any]:
        units: list[dict[str, str]] = []
        figure_count = 0
        table_count = 0
        formula_count = 0
        for candidate in candidates:
            section_id = str(candidate.get("section_id") or "")
            for block in candidate.get("paragraphs") or []:
                text = str(block or "").strip()
                if not text:
                    continue
                if "[[FIGURE]]" in text:
                    try:
                        directives = parse_figure_block(text)
                    except FigureProtocolError:
                        directives = []
                    figure_count += len(directives)
                    units.extend(
                        {"section_id": section_id, "kind": "FIGURE_CAPTION", "text": item.caption}
                        for item in directives
                        if item.caption
                    )
                    continue
                if text.startswith("[[TABLE]]"):
                    table_count += 1
                    raw = text.removeprefix("[[TABLE]]").strip()
                    for line in raw.splitlines():
                        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
                        for cell in cells:
                            if cell and not re.fullmatch(r"[-:]+", cell):
                                units.append({"section_id": section_id, "kind": "TABLE_CELL", "text": cell})
                    continue
                for marker, kind in (
                    ("[[H2]]", "HEADING"),
                    ("[[H3]]", "HEADING"),
                    ("[[FORMULA]]", "FORMULA"),
                    ("[[REFERENCE]]", "REFERENCE"),
                    ("[[BULLET]]", "LIST"),
                    ("[[NUMBER]]", "LIST"),
                ):
                    if text.startswith(marker):
                        text = text.removeprefix(marker).strip()
                        if kind == "FORMULA":
                            formula_count += 1
                        unit = {"section_id": section_id, "kind": kind, "text": text}
                        if kind == "HEADING" and units and units[-1] == unit:
                            break
                        units.append(unit)
                        break
                else:
                    units.append({"section_id": section_id, "kind": "PARAGRAPH", "text": text})
        return {
            "units": units,
            "figure_count": figure_count,
            "table_count": table_count,
            "formula_count": formula_count,
        }

    @staticmethod
    def _fuzzy_unit_present(unit: str, document: str) -> bool:
        if unit in document:
            return True
        if len(unit) < 20:
            return False
        chunk_size = 16
        chunks = [unit[index : index + chunk_size] for index in range(0, len(unit), chunk_size)]
        hits = sum(1 for chunk in chunks if len(chunk) >= 4 and chunk in document)
        return hits / max(1, len(chunks)) >= 0.8

    @staticmethod
    def _normalize(text: str) -> str:
        return re.sub(r"[^0-9A-Za-z\u4e00-\u9fff]+", "", text or "").lower()

    @staticmethod
    def _docx_visible_text(document: Document) -> str:
        values = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        for table in document.tables:
            values.extend(cell.text for row in table.rows for cell in row.cells if cell.text)
        return "\n".join(values)

    @staticmethod
    def _pdf_text(pdf_path: Path) -> str:
        # Poppler preserves the logical Chinese text layer produced by LibreOffice
        # more reliably than pypdf for CID/subset fonts.  pypdf may visually split
        # or reorder glyphs and create false content-loss findings even though the
        # rendered PDF and Poppler text are complete.
        pdftotext = shutil.which("pdftotext")
        if pdftotext:
            completed = subprocess.run(
                [pdftotext, "-layout", str(pdf_path), "-"],
                check=True,
                capture_output=True,
            )
            return completed.stdout.decode("utf-8", errors="replace")
        reader = PdfReader(str(pdf_path))
        return "\n".join((page.extract_text() or "") for page in reader.pages)

    def _paragraph_sections(self, document: Document) -> dict[int, str]:
        current = ""
        result: dict[int, str] = {}
        for index, paragraph in enumerate(document.paragraphs):
            level = self._heading_level(paragraph.style.name if paragraph.style else "")
            if level == 1 and paragraph.text.strip() != "目录":
                current = paragraph.text.strip()
            result[index] = current
        return result

    @staticmethod
    def _sha_file(path: Path) -> str:
        from .util import sha256_bytes

        return sha256_bytes(path.read_bytes())

    @staticmethod
    def _owner_counts(findings: list[dict[str, Any]]) -> dict[str, int]:
        counts: dict[str, int] = {}
        for item in findings:
            owner = str(item.get("owner") or "UNROUTED")
            counts[owner] = counts.get(owner, 0) + 1
        return counts

    @staticmethod
    def _enrich_finding(finding: dict[str, Any]) -> None:
        code = str(finding.get("code") or "").upper()
        content_codes = {"D5_REFERENCE_MISALIGNMENT"}
        if finding.get("category") == "CONTENT" or code.startswith("D5_CONTENT_") or code in content_codes:
            finding.setdefault("category", "CONTENT")
            finding.setdefault("target_type", "SECTION_CANDIDATE")
            finding["owner"] = "WRITING_AGENT"
            finding["suggested_route"] = "WRITING_AGENT"
            finding.setdefault(
                "repair_instruction",
                "由责任章节写作 Agent 定向修复，重新通过 Expression Critic 与全文 Integration Critic 后再导出。",
            )
        else:
            finding.setdefault("category", "FORMAT")
            finding.setdefault("target_type", "DOCX_PDF_DELIVERY")
            finding["owner"] = "EXPORT_ENGINEERING"
            finding["suggested_route"] = "EXPORT_ENGINEERING"
            finding.setdefault(
                "repair_instruction",
                "修复导出、渲染或版式实现；保持候选集合不变，重新导出并由新的 Delivery Validator 复核。",
            )
        finding.setdefault("repairable", True)
        finding.setdefault("evidence_refs", [])
        finding.setdefault("responsible_section_ids", [])
        if not finding.get("target_path_or_span"):
            sections = [str(item) for item in finding.get("responsible_section_ids") or [] if item]
            finding["target_path_or_span"] = (
                "candidate_sections." + ",".join(sections) if sections else str(finding.get("location") or "document")
            )

    @staticmethod
    def _finding(
        code: str,
        severity: str,
        message: str,
        *,
        location: str,
        blocking: bool,
        category: str | None = None,
        target_type: str | None = None,
        responsible_section_ids: list[str] | None = None,
        evidence: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        value = RuntimeDeliveryValidator._finding(
            code, severity, message, location=location, blocking=blocking
        )
        if category:
            value["category"] = category
        if target_type:
            value["target_type"] = target_type
        if responsible_section_ids:
            value["responsible_section_ids"] = responsible_section_ids
        if evidence:
            value["evidence"] = evidence
        return value
