from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
from copy import deepcopy
from pathlib import Path

from PIL import Image as PILImage
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .figure_protocol import (
    ARTIFACT_SCHEME,
    FigureDirective,
    FigureProtocolError,
    parse_figure_block,
    resolve_figure_reference,
)


class ExportRenderMixin:
    @staticmethod
    def _document_title(name: str) -> str:
        cleaned = re.sub(r"（智能体系统测试）|\(智能体系统测试\)", "", name).strip()
        cleaned = re.sub(r"项目申请书$", "", cleaned).strip()
        return cleaned

    def _append_block(self, document: Document, text: str, list_state: dict[str, int] | None = None) -> None:
        text = (text or "").strip()
        if not text:
            return
        if "[[FIGURE]]" in text:
            directives = parse_figure_block(text)
            if not directives:
                raise FigureProtocolError("Figure marker could not be parsed")
            for directive in directives:
                self._append_figure(document, directive)
            return
        if text.startswith("[[H2]]"):
            if list_state is not None:
                list_state["number"] = 0
            title = text.removeprefix("[[H2]]").strip()
            if document.paragraphs:
                previous = document.paragraphs[-1]
                previous_level = self._heading_level(previous)
                if previous_level == 2 and previous.text.strip() == title:
                    return
            heading = document.add_heading(title, level=2)
            if title == "参考文献" and heading.runs:
                self._set_run_font(heading.runs[0], "Noto Sans CJK SC", 13, bold=True)
            return
        if text.startswith("[[H3]]"):
            if list_state is not None:
                list_state["number"] = 0
            title = text.removeprefix("[[H3]]").strip()
            if document.paragraphs:
                previous = document.paragraphs[-1]
                previous_level = self._heading_level(previous)
                if previous_level == 3 and previous.text.strip() == title:
                    return
            document.add_heading(title, level=3)
            return
        if text.startswith("[[TABLE]]"):
            self._append_table(document, text.removeprefix("[[TABLE]]").strip())
            return
        if text.startswith("[[FORMULA]]"):
            self._append_formula(document, text.removeprefix("[[FORMULA]]").strip())
            return
        if text.startswith("[[REFERENCE]]"):
            self._append_reference(document, text.removeprefix("[[REFERENCE]]").strip())
            return
        if text.startswith("[[BULLET]]"):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.first_line_indent = None
            run = paragraph.add_run(text.removeprefix("[[BULLET]]").strip())
            self._set_run_font(run, "Noto Serif CJK SC", 11)
            return
        if text.startswith("[[NUMBER]]"):
            if list_state is None:
                list_state = {"number": 0}
            list_state["number"] = list_state.get("number", 0) + 1
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Cm(-0.65)
            paragraph.paragraph_format.left_indent = Cm(0.65)
            run = paragraph.add_run(f"{list_state['number']}. {text.removeprefix('[[NUMBER]]').strip()}")
            self._set_run_font(run, "Noto Serif CJK SC", 11)
            return
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        self._set_run_font(run, "Noto Serif CJK SC", 11)

    def _append_table(self, document: Document, raw: str) -> None:
        rows = []
        for line in raw.splitlines():
            line = line.strip().strip("|")
            if not line or set(line.replace("|", "").replace("-", "").replace(":", "")) == set():
                continue
            rows.append([cell.strip() for cell in line.split("|")])
        if not rows:
            return
        width = max(len(row) for row in rows)
        rows = [row + [""] * (width - len(row)) for row in rows]
        table = document.add_table(rows=len(rows), cols=width)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for row_index, values in enumerate(rows):
            row = table.rows[row_index]
            tr_pr = row._tr.get_or_add_trPr()
            cant_split = OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)
            if row_index == 0:
                tbl_header = OxmlElement("w:tblHeader")
                tbl_header.set(qn("w:val"), "true")
                tr_pr.append(tbl_header)
            for column_index, value in enumerate(values):
                cell = row.cells[column_index]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.text = value
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.first_line_indent = None
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        self._set_run_font(run, "Noto Serif CJK SC", 9.0, bold=row_index == 0)
        document.add_paragraph().paragraph_format.space_after = Pt(0)

    def _append_formula(self, document: Document, formula: str) -> None:
        """Render reviewed LaTeX as a native Word equation (OMML).

        The reviewed formula string is preserved in the candidate snapshot.  Pandoc
        performs only a delivery-format transformation from LaTeX math to OMML; it
        does not alter the reviewed proposal content.
        """
        if not formula:
            return
        pandoc = shutil.which("pandoc")
        if not pandoc:
            raise RuntimeError("pandoc is required to render [[FORMULA]] blocks as Word equations")

        # LibreOffice renders operator names more cleanly when Pandoc receives
        # \operatorname rather than \mathrm for multi-letter functions.
        normalized = re.sub(r"\\mathrm\{([^{}]+)\}", r"\\operatorname{\1}", formula)
        with tempfile.TemporaryDirectory(prefix="proposal-formula-") as tmp:
            tmp_dir = Path(tmp)
            source = tmp_dir / "formula.md"
            converted = tmp_dir / "formula.docx"
            source.write_text(f"$$\n{normalized}\n$$\n", encoding="utf-8")
            completed = subprocess.run(
                [pandoc, str(source), "-o", str(converted)],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if completed.returncode != 0 or not converted.is_file():
                raise RuntimeError(
                    "pandoc failed to convert proposal formula to OMML: "
                    + (completed.stderr or completed.stdout or "unknown error")
                )
            source_doc = Document(str(converted))
            if not source_doc.paragraphs:
                raise RuntimeError("pandoc produced no formula paragraph")
            source_paragraph = source_doc.paragraphs[0]._p

            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.first_line_indent = None
            for child in source_paragraph:
                if child.tag == qn("w:pPr"):
                    continue
                paragraph._p.append(deepcopy(child))

    def _append_reference(self, document: Document, reference: str) -> None:
        if not reference:
            return
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.75)
        paragraph.paragraph_format.first_line_indent = Cm(-0.75)
        run = paragraph.add_run(reference)
        self._set_run_font(run, "Noto Serif CJK SC", 9.5)

    @staticmethod
    def _set_run_font(run, font_name: str, size: float, bold: bool = False) -> None:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        run.font.size = Pt(size)
        run.bold = bold

    def _add_page_numbers(self, document: Document) -> None:
        for section in document.sections:
            paragraph = section.footer.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            begin = OxmlElement("w:fldChar")
            begin.set(qn("w:fldCharType"), "begin")
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = " PAGE "
            end = OxmlElement("w:fldChar")
            end.set(qn("w:fldCharType"), "end")
            run._r.extend([begin, instr, end])
            self._set_run_font(run, "Noto Serif CJK SC", 10)

    def _append_figure(self, document: Document, directive: FigureDirective | str) -> None:
        if isinstance(directive, str):
            parsed = parse_figure_block("[[FIGURE]]" + directive)
            if len(parsed) != 1:
                raise FigureProtocolError("Expected exactly one figure directive")
            directive = parsed[0]

        data_dir = self._figure_data_dir(directive.reference)
        img_path = resolve_figure_reference(directive.reference, data_dir)
        caption = directive.caption
        width_cm = directive.width_cm
        section = document.sections[-1]
        emu_per_cm = 360000.0
        printable_width_cm = (section.page_width - section.left_margin - section.right_margin) / emu_per_cm
        printable_height_cm = (section.page_height - section.top_margin - section.bottom_margin) / emu_per_cm
        target_width_cm = min(width_cm, printable_width_cm)
        max_figure_height_cm = min(16.5, max(8.0, printable_height_cm - 6.0))
        try:
            with PILImage.open(img_path) as image:
                pixel_width, pixel_height = image.size
                image.verify()
            target_height_cm = target_width_cm * pixel_height / max(pixel_width, 1)
        except Exception as exc:
            raise FigureProtocolError(f"Figure image cannot be decoded: {img_path.name}") from exc

        if target_height_cm > max_figure_height_cm:
            scale = max_figure_height_cm / target_height_cm
            target_width_cm *= scale
            target_height_cm = max_figure_height_cm
        # Extremely wide Mermaid diagrams otherwise become a sub-2 cm strip whose
        # labels are unreadable in DOCX/PDF. Preserve width but give the rendered
        # diagram a practical minimum height. This is a delivery-only transform; the
        # reviewed figure reference and caption remain unchanged.
        minimum_readable_height_cm = 2.8
        if target_height_cm < minimum_readable_height_cm:
            target_height_cm = min(minimum_readable_height_cm, max_figure_height_cm)

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        run = p.add_run()
        run.add_picture(str(img_path), width=Cm(target_width_cm), height=Cm(target_height_cm))
        c = document.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.keep_together = True
        crun = c.add_run(caption)
        self._set_run_font(crun, "Noto Serif CJK SC", 10)
        c.paragraph_format.space_after = Pt(6)

    def _figure_data_dir(self, reference: str) -> Path:
        settings = getattr(self, "settings", None)
        if settings is not None and getattr(settings, "data_dir", None):
            return Path(settings.data_dir)
        if reference.startswith(ARTIFACT_SCHEME):
            raise FigureProtocolError("APP_DATA_DIR is required for artifact references")
        path = Path(reference)
        return path.resolve().parent if path.is_absolute() else Path.cwd()
