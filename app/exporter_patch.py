from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from .util import sha256_bytes


class ExportPatchMixin:
    def _patch_template(self, template_path: Path, output_path: Path, candidates: list[dict[str, Any]]) -> tuple[Path, dict[str, Any]]:
        document = Document(template_path)
        before = [p.text for p in document.paragraphs]
        changed_sections: list[dict[str, Any]] = []
        skipped_sections: list[dict[str, Any]] = []
        for candidate in candidates:
            title = (candidate.get("section_title") or "").strip()
            if not title:
                skipped_sections.append({"run_id": candidate["run_id"], "reason": "missing_section_title"})
                continue
            if candidate.get("contains_complex_content"):
                skipped_sections.append({"run_id": candidate["run_id"], "section_title": title, "reason": "complex_section_requires_manual_patch"})
                continue
            result = self._replace_section(document, title, [p for p in candidate["paragraphs"] if p.strip()])
            if result:
                changed_sections.append({"run_id": candidate["run_id"], "section_title": title, **result})
            else:
                skipped_sections.append({"run_id": candidate["run_id"], "section_title": title, "reason": "heading_not_found"})
        if not changed_sections:
            return self._generate_document({"name": template_path.stem, "id": "unknown", "security_level": "INTERNAL"}, output_path, candidates, fallback_reason="No template section could be patched safely")
        document.save(output_path)
        after = [p.text for p in document.paragraphs]
        integrity = {
            "schema_version": "1.0",
            "mode": "TARGETED_TEMPLATE_PATCH",
            "template_filename": template_path.name,
            "template_sha256": sha256_bytes(template_path.read_bytes()),
            "output_sha256": sha256_bytes(output_path.read_bytes()),
            "paragraph_count_before": len(before),
            "paragraph_count_after": len(after),
            "changed_sections": changed_sections,
            "skipped_sections": skipped_sections,
            "non_target_integrity_check": self._non_target_check(before, after, changed_sections),
        }
        return output_path, integrity

    def _replace_section(self, document: Document, title: str, new_paragraphs: list[str]) -> dict[str, Any] | None:
        paragraphs = list(document.paragraphs)
        heading_index = None
        heading_level = None
        for index, paragraph in enumerate(paragraphs):
            if paragraph.text.strip() == title:
                level = self._heading_level(paragraph)
                if level is not None:
                    heading_index, heading_level = index, level
                    break
        if heading_index is None:
            return None
        end_index = len(paragraphs)
        for index in range(heading_index + 1, len(paragraphs)):
            level = self._heading_level(paragraphs[index])
            if level is not None and level <= heading_level:
                end_index = index
                break
        removed = [p.text for p in paragraphs[heading_index + 1 : end_index]]
        anchor = paragraphs[heading_index]._p
        for paragraph in paragraphs[heading_index + 1 : end_index]:
            element = paragraph._element
            element.getparent().remove(element)
        inserted = []
        for text in new_paragraphs:
            new_element = OxmlElement("w:p")
            anchor.addnext(new_element)
            new_paragraph = Paragraph(new_element, paragraphs[heading_index]._parent)
            new_paragraph.style = document.styles["Normal"]
            new_paragraph.add_run(text)
            anchor = new_element
            inserted.append(text)
        return {"heading_index_before": heading_index, "removed_paragraph_count": len(removed), "inserted_paragraph_count": len(inserted), "removed_text_hash": sha256_bytes("\n".join(removed).encode("utf-8")), "inserted_text_hash": sha256_bytes("\n".join(inserted).encode("utf-8"))}

    @staticmethod
    def _heading_level(paragraph: Paragraph) -> int | None:
        name = paragraph.style.name if paragraph.style else ""
        match = re.search(r"(?:Heading|标题)\s*(\d+)", name, flags=re.I)
        return int(match.group(1)) if match else None

