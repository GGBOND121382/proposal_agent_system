from __future__ import annotations

import json
import re
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader

from .util import new_id, safe_filename, sha256_bytes, sha256_text

ALLOWED_EXTENSIONS = {".docx", ".pdf", ".txt", ".md", ".json", ".csv"}


def _section(section_key: str, title: str, level: int, text: str, security_level: str, *, flags: dict[str, bool] | None = None) -> dict[str, Any]:
    flags = flags or {}
    sec_id = new_id("sec")
    return {
        "section_id": sec_id,
        "section_key": section_key or sec_id,
        "title": title,
        "level": level,
        "text": text,
        "text_hash": sha256_text(text),
        "block_ids": [new_id("block")],
        "contains_table": bool(flags.get("contains_table")),
        "contains_formula": bool(flags.get("contains_formula")),
        "contains_image": bool(flags.get("contains_image")),
        "contains_comment": bool(flags.get("contains_comment")),
        "contains_revision": bool(flags.get("contains_revision")),
        "security_level": security_level,
    }


def _parse_text(text: str, security_level: str) -> list[dict[str, Any]]:
    sections: list[dict[str, Any]] = []
    current_title = "全文"
    current_level = 0
    current_lines: list[str] = []
    index = 0
    heading_re = re.compile(r"^(#{1,6})\s+(.+)$")
    chinese_heading_re = re.compile(r"^([一二三四五六七八九十]+、|\d+(?:\.\d+)*[、.])\s*(.+)$")

    def flush() -> None:
        nonlocal index, current_lines
        body = "\n".join(current_lines).strip()
        # Preserve explicit headings even when their body is intentionally empty.
        # Proposal skeletons are allowed to contain only frozen section titles;
        # dropping those headings silently collapses a multi-section contract into
        # a single document and contaminates every downstream workflow.
        if body or not sections or current_title != "全文":
            index += 1
            sections.append(_section(f"section_{index}", current_title, current_level, body, security_level))
        current_lines = []

    for line in text.splitlines():
        stripped = line.strip()
        m = heading_re.match(stripped)
        m2 = chinese_heading_re.match(stripped)
        if m:
            flush()
            current_level = len(m.group(1))
            current_title = m.group(2).strip()
        elif m2:
            flush()
            current_level = max(1, m2.group(1).count(".") + 1)
            current_title = stripped
        else:
            current_lines.append(line)
    flush()
    return sections


def parse_document(filename: str, content: bytes, role: str, security_level: str) -> dict[str, Any]:
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}")
    doc_hash = sha256_bytes(content)
    sections: list[dict[str, Any]] = []
    title = Path(filename).stem

    if ext == ".docx":
        document = Document(BytesIO(content))
        current_title = "全文"
        current_level = 0
        current: list[str] = []
        contains_table = bool(document.tables)
        contains_image = bool(document.inline_shapes)
        idx = 0

        def flush() -> None:
            nonlocal idx, current
            body = "\n".join(current).strip()
            if body or not sections:
                idx += 1
                sections.append(_section(f"section_{idx}", current_title, current_level, body, security_level, flags={"contains_table": contains_table, "contains_image": contains_image}))
            current = []

        for p in document.paragraphs:
            text = p.text.strip()
            style = (p.style.name or "") if p.style else ""
            match = re.search(r"Heading\s*(\d+)|标题\s*(\d+)", style, flags=re.I)
            if match and text:
                flush()
                current_title = text
                current_level = int(match.group(1) or match.group(2) or 1)
            elif text:
                current.append(text)
        for table in document.tables:
            rows = ["\t".join(cell.text.strip() for cell in row.cells) for row in table.rows]
            if rows:
                current.append("[表格]\n" + "\n".join(rows))
        flush()
    elif ext == ".pdf":
        reader = PdfReader(BytesIO(content))
        for i, page in enumerate(reader.pages, 1):
            text = page.extract_text() or ""
            sections.append(_section(f"page_{i}", f"第{i}页", 1, text.strip(), security_level))
    else:
        text = content.decode("utf-8", errors="replace")
        if ext == ".json":
            try:
                text = json.dumps(json.loads(text), ensure_ascii=False, indent=2)
            except json.JSONDecodeError:
                pass
        sections = _parse_text(text, security_level)

    return {
        "document_id": new_id("doc"),
        "document_version_id": new_id("docv"),
        "document_role": role,
        "title": title,
        "document_hash": doc_hash,
        "authority_rank": authority_rank_for_role(role),
        "allowed_uses": allowed_uses_for_role(role),
        "prohibited_uses": prohibited_uses_for_role(role),
        "security_level": security_level,
        "sections": sections,
        "safe_filename": safe_filename(filename),
    }


def authority_rank_for_role(role: str) -> int:
    return {
        "APPLICATION_GUIDE": 95,
        "PROJECT_BRIEF": 90,
        "CURRENT_PROPOSAL": 85,
        "TECHNICAL_DESIGN": 80,
        "EVIDENCE_MATERIAL": 80,
        "TEAM_PROFILE": 75,
        "BUDGET_MATERIAL": 75,
        "REVIEW_COMMENT": 70,
        "REFERENCE_PROPOSAL": 30,
        "OTHER": 20,
    }.get(role, 20)


def allowed_uses_for_role(role: str) -> list[str]:
    if role == "REFERENCE_PROPOSAL":
        return ["STRUCTURE", "STYLE", "ARGUMENT_PATTERN"]
    return ["FACT_EXTRACTION", "PROJECT_DEFINITION", "WRITING", "CRITIC"]


def prohibited_uses_for_role(role: str) -> list[str]:
    if role == "REFERENCE_PROPOSAL":
        return ["PROJECT_FACT", "PROJECT_METRIC", "PROJECT_ACHIEVEMENT", "TECHNICAL_DESIGN_COPY"]
    return []
