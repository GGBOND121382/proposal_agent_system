from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .util import sha256_bytes, utc_now


class ExportGenerateMixin:
    def _generate_document(self, project: dict[str, Any], path: Path, candidates: list[dict[str, Any]], fallback_reason: str | None = None) -> tuple[Path, dict[str, Any]]:
        document = Document()
        self._configure_document(document)

        title = self._document_title(project.get("name", "项目申请书"))
        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.paragraph_format.space_before = Pt(90)
        title_paragraph.paragraph_format.space_after = Pt(24)
        title_run = title_paragraph.add_run(title)
        self._set_run_font(title_run, "黑体", 22, bold=True)

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run("项目申请书")
        self._set_run_font(subtitle_run, "黑体", 18, bold=True)

        description = (project.get("description") or "").strip()
        if description:
            note = document.add_paragraph()
            note.alignment = WD_ALIGN_PARAGRAPH.CENTER
            note.paragraph_format.space_before = Pt(42)
            note_run = note.add_run(description)
            self._set_run_font(note_run, "宋体", 10)

        meta = document.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.paragraph_format.space_before = Pt(80)
        meta_run = meta.add_run(f"资料等级：{project.get('security_level', 'INTERNAL')}\n生成时间：{utc_now()[:10]}")
        self._set_run_font(meta_run, "宋体", 11)
        document.add_page_break()

        if candidates:
            document.add_heading("目录", level=1)
            for candidate in candidates:
                p = document.add_paragraph()
                p.paragraph_format.first_line_indent = None
                p.paragraph_format.left_indent = Cm(0.8)
                run = p.add_run(candidate.get("section_title") or "未命名章节")
                self._set_run_font(run, "宋体", 12)
            document.add_page_break()
            for candidate in candidates:
                document.add_heading(candidate.get("section_title") or "未命名章节", level=1)
                list_state = {"number": 0}
                for block in candidate.get("paragraphs", []):
                    self._append_block(document, block, list_state)
        else:
            document.add_paragraph("当前尚无通过写作与审查流程的正文候选。")

        self._add_page_numbers(document)
        document.save(path)
        integrity = {
            "schema_version": "1.0",
            "mode": "GENERATED_DOCUMENT",
            "fallback_reason": fallback_reason,
            "output_sha256": sha256_bytes(path.read_bytes()),
            "candidate_count": len(candidates),
            "changed_sections": [{"run_id": c["run_id"], "section_title": c.get("section_title")} for c in candidates],
            "skipped_sections": [],
            "non_target_integrity_check": {"status": "NOT_APPLICABLE", "reason": "A clean multi-section document was generated; audit metadata is stored in the sidecar manifest."},
        }
        return path, integrity

    def _configure_document(self, document: Document) -> None:
        section = document.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.5)
        section.different_first_page_header_footer = True

        normal = document.styles["Normal"]
        normal.font.name = "宋体"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        normal.font.size = Pt(12)
        normal.paragraph_format.line_spacing = 1.5
        normal.paragraph_format.first_line_indent = Pt(24)
        normal.paragraph_format.space_after = Pt(6)

        for style_name, font_name, size in [("Heading 1", "黑体", 16), ("Heading 2", "黑体", 14), ("Heading 3", "楷体", 12)]:
            style = document.styles[style_name]
            style.font.name = font_name
            style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
            style.font.size = Pt(size)
            style.font.bold = True
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.keep_with_next = True

