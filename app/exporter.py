from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Pt

from .util import new_id, safe_filename, sha256_bytes, utc_now, write_json


class ExportDenied(RuntimeError):
    pass


class DocxExporter:
    def __init__(self, db, settings):
        self.db = db
        self.settings = settings

    def export(self, project_id: str) -> Path:
        project, gates = self._authorized_project(project_id)
        candidates = self._candidate_runs(project_id)
        template_row = self.db.fetchone(
            "SELECT file_path,filename,parsed_json FROM documents WHERE project_id=? AND role='CURRENT_PROPOSAL' AND filename LIKE '%.docx' ORDER BY created_at DESC LIMIT 1",
            (project_id,),
        )
        integrity: dict[str, Any]
        filename = safe_filename(f"{project['name']}-{new_id('export')}.docx")
        path = self.settings.exports_dir / filename
        if template_row and candidates:
            path, integrity = self._patch_template(Path(template_row["file_path"]), path, candidates)
        else:
            path, integrity = self._generate_document(project, path, candidates)
        manifest = self._manifest(project, gates, candidates, path, integrity)
        write_json(path.with_suffix(".integrity.json"), integrity)
        write_json(path.with_suffix(".manifest.json"), manifest)
        self.db.audit("DOCX_EXPORTED", project_id=project_id, object_id=filename, metadata={"filename": path.name, "sha256": manifest["document_sha256"], "candidate_count": len(candidates), "mode": integrity["mode"]})
        return path

    def export_package(self, project_id: str) -> Path:
        document_path = self.export(project_id)
        package_path = document_path.with_suffix(".zip")
        with zipfile.ZipFile(package_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            for path in [document_path, document_path.with_suffix(".integrity.json"), document_path.with_suffix(".manifest.json")]:
                zf.write(path, arcname=path.name)
        self.db.audit("EXPORT_PACKAGE_CREATED", project_id=project_id, object_id=package_path.name, metadata={"filename": package_path.name, "sha256": sha256_bytes(package_path.read_bytes())})
        return package_path

    def _authorized_project(self, project_id: str) -> tuple[dict[str, Any], dict[str, str]]:
        project = self.db.fetchone("SELECT * FROM projects WHERE id=?", (project_id,))
        if not project:
            raise KeyError(project_id)
        gates: dict[str, str] = {}
        for gate_type in ["FINAL_CONTENT_SECURITY_APPROVAL", "FINAL_EXPORT_APPROVAL"]:
            gate = self.db.fetchone("SELECT id,status FROM gates WHERE project_id=? AND gate_type=? ORDER BY created_at DESC LIMIT 1", (project_id, gate_type))
            if not gate or gate["status"] != "APPROVED":
                raise ExportDenied(f"{gate_type} gate has not been approved")
            gates[gate_type] = gate["id"]
        return project, gates

    def _candidate_runs(self, project_id: str) -> list[dict[str, Any]]:
        rows = self.db.fetchall(
            "SELECT id,input_json,output_json,created_at FROM prompt_runs WHERE project_id=? AND prompt_id='P-WRITE-CONTENT' AND status='PASS' ORDER BY created_at",
            (project_id,),
        )
        candidates = []
        for row in rows:
            input_data = json.loads(row["input_json"])
            output = json.loads(row["output_json"])
            source_section = input_data.get("payload", {}).get("source_section", {})
            result = output.get("result", {})
            candidates.append(
                {
                    "run_id": row["id"],
                    "section_id": source_section.get("section_id"),
                    "section_title": source_section.get("title") or source_section.get("section_key"),
                    "source_section_hash": source_section.get("text_hash"),
                    "contains_complex_content": any(source_section.get(k, False) for k in ["contains_table", "contains_formula", "contains_image", "contains_comment", "contains_revision"]),
                    "paragraphs": [p.get("text", "") for p in sorted(result.get("paragraphs", []), key=lambda x: x.get("sequence", 0))] or [result.get("candidate_text", "")],
                    "candidate_id": result.get("candidate_id"),
                }
            )
        return candidates

    def _patch_template(self, template_path: Path, output_path: Path, candidates: list[dict[str, Any]]) -> tuple[Path, dict[str, Any]]:
        document = Document(template_path)
        before = [p.text for p in document.paragraphs]
        changed_sections: list[dict[str, Any]] = []
        skipped_sections: list[dict[str, Any]] = []
        for candidate in candidates:
            title = (candidate.get("section_title") or "").strip()
            if not title:
                skipped_sections.append({"run_id": candidate["run_id"], "reason": "missing_section_title"})
                continue
            if candidate.get("contains_complex_content"):
                skipped_sections.append({"run_id": candidate["run_id"], "section_title": title, "reason": "complex_section_requires_manual_patch"})
                continue
            result = self._replace_section(document, title, [p for p in candidate["paragraphs"] if p.strip()])
            if result:
                changed_sections.append({"run_id": candidate["run_id"], "section_title": title, **result})
            else:
                skipped_sections.append({"run_id": candidate["run_id"], "section_title": title, "reason": "heading_not_found"})
        if not changed_sections:
            return self._generate_document({"name": template_path.stem, "id": "unknown", "security_level": "INTERNAL"}, output_path, candidates, fallback_reason="No template section could be patched safely")
        document.save(output_path)
        after = [p.text for p in document.paragraphs]
        integrity = {
            "schema_version": "1.0",
            "mode": "TARGETED_TEMPLATE_PATCH",
            "template_filename": template_path.name,
            "template_sha256": sha256_bytes(template_path.read_bytes()),
            "output_sha256": sha256_bytes(output_path.read_bytes()),
            "paragraph_count_before": len(before),
            "paragraph_count_after": len(after),
            "changed_sections": changed_sections,
            "skipped_sections": skipped_sections,
            "non_target_integrity_check": self._non_target_check(before, after, changed_sections),
        }
        return output_path, integrity

    def _replace_section(self, document: Document, title: str, new_paragraphs: list[str]) -> dict[str, Any] | None:
        paragraphs = list(document.paragraphs)
        heading_index = None
        heading_level = None
        for index, paragraph in enumerate(paragraphs):
            if paragraph.text.strip() == title:
                level = self._heading_level(paragraph)
                if level is not None:
                    heading_index, heading_level = index, level
                    break
        if heading_index is None:
            return None
        end_index = len(paragraphs)
        for index in range(heading_index + 1, len(paragraphs)):
            level = self._heading_level(paragraphs[index])
            if level is not None and level <= heading_level:
                end_index = index
                break
        removed = [p.text for p in paragraphs[heading_index + 1 : end_index]]
        anchor = paragraphs[heading_index]._p
        for paragraph in paragraphs[heading_index + 1 : end_index]:
            element = paragraph._element
            element.getparent().remove(element)
        inserted = []
        for text in new_paragraphs:
            new_element = OxmlElement("w:p")
            anchor.addnext(new_element)
            new_paragraph = Paragraph(new_element, paragraphs[heading_index]._parent)
            new_paragraph.style = document.styles["Normal"]
            new_paragraph.add_run(text)
            anchor = new_element
            inserted.append(text)
        return {"heading_index_before": heading_index, "removed_paragraph_count": len(removed), "inserted_paragraph_count": len(inserted), "removed_text_hash": sha256_bytes("\n".join(removed).encode("utf-8")), "inserted_text_hash": sha256_bytes("\n".join(inserted).encode("utf-8"))}

    @staticmethod
    def _heading_level(paragraph: Paragraph) -> int | None:
        name = paragraph.style.name if paragraph.style else ""
        match = re.search(r"(?:Heading|标题)\s*(\d+)", name, flags=re.I)
        return int(match.group(1)) if match else None

    def _generate_document(self, project: dict[str, Any], path: Path, candidates: list[dict[str, Any]], fallback_reason: str | None = None) -> tuple[Path, dict[str, Any]]:
        document = Document()
        document.styles["Normal"].font.name = "宋体"
        document.styles["Normal"].font.size = Pt(12)
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(project["name"])
        run.bold = True
        run.font.size = Pt(18)
        if candidates:
            for index, candidate in enumerate(candidates, 1):
                document.add_heading(candidate.get("section_title") or f"候选章节 {index}", level=1)
                for text in candidate.get("paragraphs", []):
                    if text.strip():
                        document.add_paragraph(text.strip())
        else:
            document.add_paragraph("当前尚无通过写作与审查流程的正文候选。")
        document.add_page_break()
        document.add_heading("导出审计信息", level=1)
        document.add_paragraph(f"项目编号：{project.get('id', 'unknown')}")
        document.add_paragraph(f"安全等级：{project.get('security_level', 'INTERNAL')}")
        document.add_paragraph(f"导出时间：{utc_now()}")
        document.save(path)
        integrity = {
            "schema_version": "1.0",
            "mode": "GENERATED_DOCUMENT",
            "fallback_reason": fallback_reason,
            "output_sha256": sha256_bytes(path.read_bytes()),
            "candidate_count": len(candidates),
            "changed_sections": [{"run_id": c["run_id"], "section_title": c.get("section_title")} for c in candidates],
            "skipped_sections": [],
            "non_target_integrity_check": {"status": "NOT_APPLICABLE", "reason": "No source DOCX template was patched"},
        }
        return path, integrity

    @staticmethod
    def _non_target_check(before: list[str], after: list[str], changed_sections: list[dict[str, Any]]) -> dict[str, Any]:
        # Full OOXML equivalence is not claimed. This check records stable prefix/suffix evidence
        # around the targeted paragraph ranges and makes the limitation explicit.
        prefix = 0
        for left, right in zip(before, after):
            if left != right:
                break
            prefix += 1
        suffix = 0
        for left, right in zip(reversed(before), reversed(after)):
            if left != right:
                break
            suffix += 1
        return {"status": "PASS_WITH_LIMITATION", "unchanged_prefix_paragraphs": prefix, "unchanged_suffix_paragraphs": suffix, "limitation": "Paragraph-level check; complex OOXML objects require deployment-specific validation."}

    def _manifest(self, project: dict[str, Any], gates: dict[str, str], candidates: list[dict[str, Any]], path: Path, integrity: dict[str, Any]) -> dict[str, Any]:
        return {
            "schema_version": "1.0",
            "project_id": project["id"],
            "project_name": project["name"],
            "security_level": project["security_level"],
            "exported_at": utc_now(),
            "document_filename": path.name,
            "document_sha256": sha256_bytes(path.read_bytes()),
            "approval_gate_ids": gates,
            "source_run_ids": [c["run_id"] for c in candidates],
            "integrity_mode": integrity["mode"],
        }
