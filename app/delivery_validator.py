from __future__ import annotations

import re
import shutil
import subprocess
from pathlib import Path
from typing import Any

from PIL import Image
from docx import Document
from pypdf import PdfReader

from .util import sha256_bytes, utc_now, write_json


class DeliveryValidationError(RuntimeError):
    pass


class DeliveryValidator:
    """Deterministic structural and page-visual validation for final deliveries."""

    PLACEHOLDER_PATTERNS = {
        "PLACEHOLDER_BRACES": re.compile(r"\{\{[^{}]+\}\}|<<[^<>]+>>"),
        "PLACEHOLDER_WORD": re.compile(
            r"(?im)\b(?:TODO|TBD|PLACEHOLDER|FIXME)\b|"
            r"(?:^|[\s【\[])待补充(?:项|内容|材料|数据)?(?:[\s】\]]*$|[:：])|"
            r"此处填写|示例文本"
        ),
    }
    INTERNAL_PATTERNS = {
        "INTERNAL_RUNTIME_TERM": re.compile(
            r"\b(?:PROMPT_TRACE|PROMPT_OUTPUT|SIMULATED|REPLAY|MOCK|workflow_id|run_id|skill_run|APP_DATA_DIR)\b",
            re.I,
        ),
        "INTERNAL_PATH": re.compile(
            r"(?:[A-Za-z]:\\(?:[^\s]+\\)+[^\s]*|/(?:tmp|home|var|mnt|app)/[^\s]+|artifact://[^\s]+)"
        ),
        "UNRENDERED_DIRECTIVE": re.compile(r"\[\[(?:FIGURE|TABLE|FORMULA|H2|H3|BULLET|NUMBER|REFERENCE)\]\]"),
    }

    def __init__(self, settings):
        self.settings = settings

    def validate(
        self,
        docx_path: Path,
        pdf_path: Path,
        *,
        expected_sections: list[str] | None = None,
        screenshots_dir: Path | None = None,
    ) -> dict[str, Any]:
        docx_path = docx_path.resolve()
        pdf_path = pdf_path.resolve()
        screenshots_dir = screenshots_dir or docx_path.parent / f"{docx_path.stem}-pages"
        structure = self.validate_structure(
            docx_path, pdf_path, expected_sections=expected_sections or []
        )
        visual = self.validate_visual(docx_path, pdf_path, screenshots_dir=screenshots_dir)
        findings = structure["findings"] + visual["findings"]
        blocking = [item for item in findings if item.get("blocking", True)]
        report = {
            "schema_version": "1.0",
            "validated_at": utc_now(),
            "status": "FAIL" if blocking else "PASS",
            "docx_filename": docx_path.name,
            "docx_sha256": sha256_bytes(docx_path.read_bytes()),
            "pdf_filename": pdf_path.name,
            "pdf_sha256": sha256_bytes(pdf_path.read_bytes()),
            "structure_report": structure["report_path"],
            "visual_report": visual["report_path"],
            "screenshot_dir": str(screenshots_dir),
            "finding_count": len(findings),
            "blocking_finding_count": len(blocking),
            "findings": findings,
        }
        overall_path = docx_path.with_suffix(".delivery-validation.json")
        write_json(overall_path, report)
        report["report_path"] = str(overall_path)
        return report

    def validate_structure(
        self,
        docx_path: Path,
        pdf_path: Path,
        *,
        expected_sections: list[str],
    ) -> dict[str, Any]:
        findings: list[dict[str, Any]] = []
        document = Document(str(docx_path))
        paragraphs = list(document.paragraphs)
        visible_texts = [p.text.strip() for p in paragraphs if p.text.strip()]
        visible_text = "\n".join(visible_texts)

        headings: list[tuple[int, str, int]] = []
        for index, paragraph in enumerate(paragraphs):
            title = paragraph.text.strip()
            level = self._heading_level(paragraph.style.name if paragraph.style else "")
            if title and level:
                headings.append((index, title, level))
        heading_titles = [title for _, title, _ in headings]

        for expected in expected_sections:
            if expected and expected not in heading_titles:
                findings.append(self._finding(
                    "D5_MISSING_SECTION", "P0", f"缺少预期章节：{expected}",
                    location="DOCX", blocking=True,
                ))

        # A parent heading may be followed immediately by a lower-level heading, and a
        # subsection may contain a table or figure without an ordinary text paragraph.
        # Inspect the real DOCX XML block range until the next heading of the same or a
        # higher level instead of looking only at paragraphs before the next heading.
        body_children = list(document.element.body.iterchildren())
        child_index_by_id = {id(child): index for index, child in enumerate(body_children)}
        for position, (paragraph_index, title, level) in enumerate(headings):
            if title == "目录":
                continue
            next_boundary = len(paragraphs)
            for next_paragraph_index, _next_title, next_level in headings[position + 1 :]:
                if next_level <= level:
                    next_boundary = next_paragraph_index
                    break
            start_child = child_index_by_id.get(id(paragraphs[paragraph_index]._p), -1)
            if next_boundary < len(paragraphs):
                end_child = child_index_by_id.get(id(paragraphs[next_boundary]._p), len(body_children))
            else:
                end_child = len(body_children)
            has_content = False
            for child in body_children[start_child + 1 : end_child]:
                local_name = child.tag.rsplit("}", 1)[-1]
                if local_name == "tbl":
                    has_content = True
                    break
                if local_name != "p":
                    continue
                paragraph = next((item for item in paragraphs if item._p is child), None)
                if paragraph is None:
                    continue
                child_level = self._heading_level(paragraph.style.name if paragraph.style else "")
                if child_level:
                    continue
                if paragraph.text.strip() or child.xpath('.//w:drawing') or child.xpath('.//w:object'):
                    has_content = True
                    break
            if not has_content and level <= 3:
                findings.append(self._finding(
                    "D5_EMPTY_SECTION", "P1", f"章节没有正文：{title}",
                    location=f"DOCX paragraph {paragraph_index + 1}", blocking=True,
                ))

        duplicates = sorted({title for title in heading_titles if heading_titles.count(title) > 1 and title != "目录"})
        for title in duplicates:
            findings.append(self._finding(
                "D5_DUPLICATE_HEADING", "P1", f"章节标题重复：{title}",
                location="DOCX", blocking=True,
            ))

        for code, pattern in {**self.PLACEHOLDER_PATTERNS, **self.INTERNAL_PATTERNS}.items():
            for match in pattern.finditer(visible_text):
                severity = "P0" if code in {"INTERNAL_PATH", "UNRENDERED_DIRECTIVE"} else "P1"
                findings.append(self._finding(
                    f"D5_{code}", severity,
                    f"正文中发现不应进入交付物的内容：{match.group(0)[:120]}",
                    location="DOCX body", blocking=True,
                ))

        for table_index, table in enumerate(document.tables, 1):
            rows = len(table.rows)
            columns = len(table.columns)
            if rows < 2 or columns < 1:
                findings.append(self._finding(
                    "D5_ABNORMAL_TABLE", "P1",
                    f"表格 {table_index} 的行列结构异常（{rows}×{columns}）",
                    location=f"DOCX table {table_index}", blocking=True,
                ))
                continue
            if columns > 12:
                findings.append(self._finding(
                    "D5_TABLE_TOO_WIDE", "P1",
                    f"表格 {table_index} 包含 {columns} 列，可能不可读",
                    location=f"DOCX table {table_index}", blocking=True,
                ))
            header = [cell.text.strip() for cell in table.rows[0].cells]
            if not any(header) or any(not value for value in header):
                findings.append(self._finding(
                    "D5_TABLE_HEADER_EMPTY", "P1",
                    f"表格 {table_index} 存在空表头",
                    location=f"DOCX table {table_index}", blocking=True,
                ))
            for row_index, row in enumerate(table.rows, 1):
                if not any(cell.text.strip() for cell in row.cells):
                    findings.append(self._finding(
                        "D5_EMPTY_TABLE_ROW", "P1",
                        f"表格 {table_index} 第 {row_index} 行为空",
                        location=f"DOCX table {table_index}", blocking=True,
                    ))

        citation_numbers = {int(value) for value in re.findall(r"\[(\d{1,3})\]", visible_text)}
        reference_numbers = {
            int(value)
            for value in re.findall(r"(?m)^\s*\[(\d{1,3})\]\s+", visible_text)
        }
        missing_references = sorted(citation_numbers - reference_numbers)
        if citation_numbers and missing_references:
            findings.append(self._finding(
                "D5_REFERENCE_MISALIGNMENT", "P1",
                "正文引用缺少对应参考文献条目：" + ", ".join(map(str, missing_references[:20])),
                location="DOCX references", blocking=True,
            ))

        try:
            reader = PdfReader(str(pdf_path))
        except Exception as exc:
            findings.append(self._finding(
                "D5_PDF_UNREADABLE", "P0", f"PDF 无法读取：{exc}",
                location="PDF", blocking=True,
            ))
            page_count = 0
        else:
            page_count = len(reader.pages)
            if page_count < 1:
                findings.append(self._finding(
                    "D5_PDF_EMPTY", "P0", "PDF 没有页面", location="PDF", blocking=True,
                ))
            for page_number, page in enumerate(reader.pages, 1):
                try:
                    text = page.extract_text() or ""
                except Exception as exc:
                    findings.append(self._finding(
                        "D5_PDF_TEXT_EXTRACTION_FAILED", "P1",
                        f"第 {page_number} 页文本提取失败：{exc}",
                        location=f"PDF page {page_number}", blocking=True,
                    ))
                    continue
                if "�" in text or "[[FORMULA]]" in text or "[[FIGURE]]" in text:
                    findings.append(self._finding(
                        "D5_PDF_RENDERING_ARTIFACT", "P1",
                        f"第 {page_number} 页包含替换字符或未渲染指令",
                        location=f"PDF page {page_number}", blocking=True,
                    ))

        report = {
            "schema_version": "1.0",
            "validated_at": utc_now(),
            "status": "FAIL" if any(item["blocking"] for item in findings) else "PASS",
            "docx": {
                "paragraph_count": len(paragraphs),
                "heading_count": len(headings),
                "table_count": len(document.tables),
                "figure_count": len(document.inline_shapes),
                "expected_sections": expected_sections,
                "actual_headings": heading_titles,
            },
            "pdf": {"page_count": page_count},
            "findings": findings,
        }
        path = docx_path.with_suffix(".structure-findings.json")
        write_json(path, report)
        report["report_path"] = str(path)
        return report

    def validate_visual(
        self,
        docx_path: Path,
        pdf_path: Path,
        *,
        screenshots_dir: Path,
    ) -> dict[str, Any]:
        findings: list[dict[str, Any]] = []
        screenshots_dir.mkdir(parents=True, exist_ok=True)
        for old in screenshots_dir.glob("page-*.png"):
            old.unlink()

        rasterizer = shutil.which("pdftoppm")
        if not rasterizer:
            raise DeliveryValidationError(
                "pdftoppm is unavailable; page visual validation cannot be silently skipped"
            )
        prefix = screenshots_dir / "page"
        completed = subprocess.run(
            [rasterizer, "-png", "-r", "120", str(pdf_path), str(prefix)],
            text=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=300,
            check=False,
        )
        if completed.returncode != 0:
            raise DeliveryValidationError(
                "PDF page rasterization failed: "
                + (completed.stderr.strip() or completed.stdout.strip())[-1200:]
            )

        pages = sorted(screenshots_dir.glob("page-*.png"))
        if not pages:
            raise DeliveryValidationError("Page rasterization produced no screenshots")

        screenshot_records: list[dict[str, Any]] = []
        for page_number, image_path in enumerate(pages, 1):
            with Image.open(image_path) as image:
                gray = image.convert("L")
                width, height = gray.size
                mask = gray.point(lambda value: 0 if value > 245 else 255)
                bbox = mask.getbbox()
                if bbox is None:
                    findings.append(self._finding(
                        "D6_BLANK_PAGE", "P1", f"第 {page_number} 页为空白页",
                        location=f"PDF page {page_number}", blocking=True,
                    ))
                    content_ratio = 0.0
                    edge_margin = None
                else:
                    left, top, right, bottom = bbox
                    bbox_area = max(1, (right - left) * (bottom - top))
                    content_ratio = bbox_area / max(1, width * height)
                    edge_margin = min(left, top, width - right, height - bottom)
                    if edge_margin <= 2:
                        findings.append(self._finding(
                            "D6_CROP_RISK", "P1",
                            f"第 {page_number} 页内容触及页面边缘，存在裁切风险",
                            location=f"PDF page {page_number}", blocking=True,
                        ))
                    if content_ratio < 0.025:
                        findings.append(self._finding(
                            "D6_ABNORMAL_WHITESPACE", "P1",
                            f"第 {page_number} 页有效内容区域过小（{content_ratio:.1%}）",
                            location=f"PDF page {page_number}", blocking=True,
                        ))
                screenshot_records.append({
                    "page": page_number,
                    "path": str(image_path),
                    "sha256": sha256_bytes(image_path.read_bytes()),
                    "width_px": width,
                    "height_px": height,
                    "content_bbox_ratio": round(content_ratio, 6),
                    "minimum_edge_margin_px": edge_margin,
                })

        document = Document(str(docx_path))
        for index, shape in enumerate(document.inline_shapes, 1):
            width_cm = shape.width.cm
            height_cm = shape.height.cm
            if width_cm < 4.0 or height_cm < 2.0:
                findings.append(self._finding(
                    "D6_FIGURE_TOO_SMALL", "P1",
                    f"图 {index} 尺寸过小（{width_cm:.2f}×{height_cm:.2f} cm）",
                    location=f"DOCX figure {index}", blocking=True,
                ))

        for table_index, table in enumerate(document.tables, 1):
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.font.size is not None and run.font.size.pt < 7:
                                findings.append(self._finding(
                                    "D6_TABLE_TEXT_TOO_SMALL", "P1",
                                    f"表格 {table_index} 存在小于 7 磅的文字",
                                    location=f"DOCX table {table_index}", blocking=True,
                                ))
                                break

        findings.extend(self._pdf_overlap_findings(pdf_path))

        report = {
            "schema_version": "1.0",
            "validated_at": utc_now(),
            "status": "FAIL" if any(item["blocking"] for item in findings) else "PASS",
            "page_count": len(pages),
            "screenshots": screenshot_records,
            "findings": findings,
        }
        path = docx_path.with_suffix(".visual-findings.json")
        write_json(path, report)
        report["report_path"] = str(path)
        return report

    def require_pass(self, report: dict[str, Any]) -> None:
        if report.get("status") != "PASS":
            codes = [item.get("code", "UNKNOWN") for item in report.get("findings", []) if item.get("blocking", True)]
            raise DeliveryValidationError(
                "Delivery validation failed: " + ", ".join(codes[:20])
            )

    def _pdf_overlap_findings(self, pdf_path: Path) -> list[dict[str, Any]]:
        findings: list[dict[str, Any]] = []
        reader = PdfReader(str(pdf_path))
        for page_number, page in enumerate(reader.pages, 1):
            boxes: list[tuple[float, float, float, float, str]] = []

            def visitor(text, cm, tm, font_dict, font_size):  # type: ignore[no-untyped-def]
                value = (text or "").strip()
                if not value or not font_size:
                    return
                x = float(tm[4])
                y = float(tm[5])
                height = max(4.0, float(font_size))
                width = max(height * 0.45, min(len(value), 80) * height * 0.48)
                boxes.append((x, y, x + width, y + height, value[:80]))

            try:
                page.extract_text(visitor_text=visitor)
            except Exception:
                continue
            overlaps = 0
            for index, left in enumerate(boxes):
                for right in boxes[index + 1 :]:
                    if left[4] == right[4]:
                        continue
                    intersection = self._intersection_area(left[:4], right[:4])
                    if intersection <= 0:
                        continue
                    left_area = max(1.0, (left[2] - left[0]) * (left[3] - left[1]))
                    right_area = max(1.0, (right[2] - right[0]) * (right[3] - right[1]))
                    if intersection / min(left_area, right_area) >= 0.65:
                        overlaps += 1
                        if overlaps >= 3:
                            break
                if overlaps >= 3:
                    break
            if overlaps >= 3:
                findings.append(self._finding(
                    "D6_TEXT_OVERLAP_RISK", "P1",
                    f"第 {page_number} 页检测到多个文本框高度重叠，存在叠字风险",
                    location=f"PDF page {page_number}", blocking=True,
                ))
        return findings

    @staticmethod
    def _intersection_area(left: tuple[float, ...], right: tuple[float, ...]) -> float:
        x1 = max(left[0], right[0])
        y1 = max(left[1], right[1])
        x2 = min(left[2], right[2])
        y2 = min(left[3], right[3])
        return max(0.0, x2 - x1) * max(0.0, y2 - y1)

    @staticmethod
    def _heading_level(style_name: str) -> int | None:
        match = re.search(r"(?:Heading|标题)\s*(\d+)", style_name or "", flags=re.I)
        return int(match.group(1)) if match else None

    @staticmethod
    def _finding(
        code: str,
        severity: str,
        message: str,
        *,
        location: str,
        blocking: bool,
    ) -> dict[str, Any]:
        return {
            "code": code,
            "severity": severity,
            "message": message,
            "location": location,
            "blocking": blocking,
            "owner": "EXPORT_ENGINEERING",
        }
