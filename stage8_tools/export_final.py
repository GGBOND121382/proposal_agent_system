from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from matplotlib.font_manager import FontProperties
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from pypdf import PdfReader


FONT_SERIF = "Noto Serif CJK SC"
FONT_SANS = "Noto Sans CJK SC"
MATPLOTLIB_SANS = FontProperties(fname="/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc")
MATPLOTLIB_SERIF = FontProperties(fname="/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc")

CHAPTER_NUMBERS = {
    "摘要与项目概览": "一",
    "研究背景与意义": "二",
    "相关工作与现有差距": "三",
    "核心概念与理论框架": "四",
    "研究问题与目标": "五",
    "研究内容与工作包": "六",
    "关键科学与技术问题": "七",
    "技术路线与人机协同流程": "八",
    "创新点": "九",
    "评价指标与验证方案": "十",
    "研究基础与条件": "十一",
    "研究计划与里程碑": "十二",
    "风险与组织管理": "十三",
    "结论": "十四",
}

TABLE_SPECS = {
    "TAB-01": {
        "caption": "表1 相关工作机制、适用范围、局限与本项目比较维度",
        "headers": ["研究方向", "代表机制", "主要适用范围", "本项目比较维度"],
        "rows": [
            ["共享状态与协作", "角色协议、对话共享、共享心智模型", "任务分工与沟通一致", "对象粒度、版本依赖、失效传播、依据追溯"],
            ["有限时间协作", "多智能体辩论、共识形成、实时决策评价", "答案改进与计算时延评价", "统一调度推理、求解、人工注意力与停止时机"],
            ["人机互补", "能力感知分工、学习转交、解释支持", "分类、推荐与单次判断", "多阶段权限交接、决定适用范围、收益归因"],
        ],
    },
    "TAB-02": {
        "caption": "表2 研究问题、目标、研究内容、验证证据与反证条件",
        "headers": ["研究问题", "对应目标与内容", "主要验证证据", "反证条件"],
        "rows": [
            ["共享状态一致性", "目标一、目标二；共享态势与增量传播", "版本一致性、传播范围、依据追溯", "不优于全量同步或人工核对"],
            ["有限窗口调度", "目标二、目标三；候选竞争与冲刺调度", "首个可行解、质量-时间曲线、按时完成率", "不优于固定轮次或最佳非冲刺基线"],
            ["权限与收益证据", "目标三、目标四；人工交接与证据归因", "人工投入、采纳、覆盖与追溯", "不能降低负担或保持最终控制"],
        ],
    },
    "TAB-03": {
        "caption": "表3 研究内容、工作包、方法、交付与验证证据对应关系",
        "headers": ["工作包", "核心内容", "主要方法", "交付与验证"],
        "rows": [
            ["共享态势建模", "对象、版本、依赖与影响子图", "统一对象关系建模；增量传播", "态势图、版本协议、一致性与传播试验"],
            ["候选生成与批判", "角色协作、候选竞争、问题路由", "角色分离编排；确定性校验", "角色合同、候选协议、可行化与局部修复试验"],
            ["人工偏好与交接", "高价值提问、确认、否决与覆盖", "偏好学习；主动交互", "偏好约束、权限记录、人工负担试验"],
            ["资源调度与停止", "动作价值、资源分配、停止判断", "边际增益调度；随时可用优化", "调度协议、同预算质量-时间对照"],
            ["原型与跨场景验证", "统一接入、对照消融、失败归因", "全过程证据评价", "原型、复现脚本、报告与失败案例库"],
        ],
    },
    "TAB-04": {
        "caption": "表4 最近工作、机制局限、新增机制、比较维度与验证证据",
        "headers": ["方向", "现有局限", "待验证机制", "验证与反证"],
        "rows": [
            ["共享态势", "消息或步骤级共享，版本依赖不足", "版本化对象与增量失效传播", "比较冲突、修复时延与重算范围"],
            ["协同调度", "固定轮次或只调度模型推理", "推理、求解与人工注意力联合调度", "比较单位时间增益、按时完成率与停止损失"],
            ["责任归因", "偏单次判断，缺少多阶段证据链", "决定范围、撤销关系与全过程归因", "比较复核效率、追溯覆盖与控制权"],
        ],
    },
    "TAB-05": {
        "caption": "表5 基线、实验组、指标和统计协议",
        "headers": ["实验模块", "对照或消融", "核心指标", "统计要求"],
        "rows": [
            ["共享状态", "无统一状态、全量同步、增量传播", "一致性、传播时延、重算范围", "配对差值、置信区间、失败案例"],
            ["候选可行化", "移除批判、工具校验或局部修复", "可行化速度、返工范围、最终质量", "报告超时、不可行与回退原因"],
            ["资源调度", "固定轮次、风险优先、贪心与完整调度", "质量-时间曲线、按时完成率", "同预算、同端点、同求解资源"],
            ["人工交互", "全程确认、最终审批、固定与主动提问", "人工时间、有效操作、采纳", "参与者和任务分层分析"],
            ["跨场景验证", "四类场景的完整方法、基线与扰动", "稳定性、追溯、跨场景保持", "中位数、效应量和自助法区间"],
        ],
    },
    "TAB-06": {
        "caption": "表6 研究基础证据等级、能力映射、证明边界与待补条件",
        "headers": ["证据类别", "可支撑能力", "证明边界", "待补材料"],
        "rows": [
            ["申请人经历说明", "优化、增量计算、复杂方案筹划", "证明相关经历与启动能力", "论文、项目任务书、成果证明"],
            ["阶段运行记录", "结构化接口、审查、局部修改与恢复", "证明智能体工程流程能力", "目标原型与真实场景测试"],
            ["计划建设条件", "参与者试验、跨场景评价", "目前尚无实测结论", "团队、数据授权、环境与统计方案"],
        ],
    },
    "TAB-07": {
        "caption": "表7 工作包依赖、阶段验收条件与验证证据",
        "headers": ["阶段", "前置依赖", "阶段工件", "验收证据"],
        "rows": [
            ["共享态势", "场景对象与规则可整理", "对象规范、版本与依赖索引", "一致性和影响传播记录"],
            ["候选协作", "共享态势稳定", "角色合同、候选集和问题清单", "可行化、批判与局部修复试验"],
            ["人工交接", "候选差异可解释", "偏好约束与权限记录", "提问价值、覆盖和撤销试验"],
            ["冲刺调度", "动作成本和证据状态可记录", "调度与停止协议", "固定策略对照和调整记录"],
            ["集成验证", "前述工件可复现", "原型、数据和分析报告", "跨场景对照、消融及失败案例"],
        ],
    },
    "TAB-08": {
        "caption": "表8 风险信号、监测、调整与责任原则",
        "headers": ["风险类别", "主要信号", "处理方式", "责任原则"],
        "rows": [
            ["状态与依赖", "版本冲突、失效未传播", "扩大重算、全量核对、恢复确认版本", "状态维护与业务确认分离"],
            ["调度估计", "收益偏差、动作反转、耗时失控", "采用确定性顺序或缩小搜索", "调度建议不替代最终决定"],
            ["人工偏好", "偏好冲突、决定范围不清", "暂停推进、澄清范围、人员确认", "授权人员保留最终确认"],
            ["跨场景评价", "收益不稳定、负担上升", "限制外推、补充样本、报告边界", "独立审查评价证据"],
        ],
    },
}

FIGURE_SPECS = {
    "FIG-01": {
        "caption": "图1 共享态势、有限窗口调度与权限交接的统一闭环框架",
        "nodes": ["共享决策态势\n对象·版本·依赖", "有限窗口调度\n生成·校验·求解", "权限交接\n偏好·确认·覆盖"],
        "footer": "候选版本与全过程证据贯穿闭环",
    },
    "FIG-02": {
        "caption": "图2 输入治理、共享态势、候选竞争、确定性校验与输出的端到端技术路线",
        "nodes": ["输入治理", "共享态势", "候选竞争", "确定性校验", "人工确认", "输出与记录"],
        "footer": "状态变化沿依赖关系触发局部更新",
    },
    "FIG-03": {
        "caption": "图3 人工交接、停止调整、版本恢复与当前最好可行方案保留流程",
        "nodes": ["候选评估", "是否需人工判断", "限定问题与证据", "人工决定", "继续改进或停止", "保留确认版本"],
        "footer": "异常时缩小搜索、切换工具或恢复至最近确认版本",
    },
    "FIG-04": {
        "caption": "图4 跨场景对照、消融、扰动与反证验证框架",
        "nodes": ["四类通用场景", "四类基线", "五组实验", "八项指标", "统计与失败案例", "反证与边界修正"],
        "footer": "相同任务快照、模型端点、求解资源和决策窗口",
    },
}


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=60, start=80, bottom=60, end=80) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def cant_split_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_run_font(run, size=10.0, bold=False, sans=False, color=None) -> None:
    font = FONT_SANS if sans else FONT_SERIF
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    set_run_font(run, 9)


def restart_page_number(section, start=1) -> None:
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:start"), str(start))


def draw_flow_figure(spec: dict, out_path: Path) -> None:
    n = len(spec["nodes"])
    fig_w = 10.6
    fig_h = 2.45 if n <= 3 else 2.3
    fig, ax = plt.subplots(figsize=(fig_w, fig_h), dpi=180)
    ax.set_xlim(0, n)
    ax.set_ylim(0, 1)
    ax.axis("off")
    for i, label in enumerate(spec["nodes"]):
        x = i + 0.5
        box = FancyBboxPatch((x - 0.38, 0.36), 0.76, 0.34,
                             boxstyle="round,pad=0.03,rounding_size=0.03",
                             linewidth=1.0, edgecolor="black", facecolor="white")
        ax.add_patch(box)
        ax.text(x, 0.53, label, ha="center", va="center", fontsize=8.2,
                fontproperties=MATPLOTLIB_SANS)
        if i < n - 1:
            arr = FancyArrowPatch((x + 0.39, 0.53), (x + 0.61, 0.53),
                                  arrowstyle="-|>", mutation_scale=10,
                                  linewidth=1.0, color="black")
            ax.add_patch(arr)
    if n == 3:
        arr = FancyArrowPatch((2.5, 0.34), (0.5, 0.34), connectionstyle="arc3,rad=0.28",
                              arrowstyle="-|>", mutation_scale=10, linewidth=0.9, color="black")
        ax.add_patch(arr)
    ax.text(n / 2, 0.13, spec["footer"], ha="center", va="center", fontsize=8,
            fontproperties=MATPLOTLIB_SERIF)
    plt.tight_layout(pad=0.15)
    fig.savefig(out_path, bbox_inches="tight", pad_inches=0.04)
    plt.close(fig)


def configure_styles(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.9)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.15)
    sec.right_margin = Cm(2.05)
    sec.header_distance = Cm(0.8)
    sec.footer_distance = Cm(0.8)
    sec.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = FONT_SERIF
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SERIF)
    normal.font.size = Pt(10.5)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(18.0)
    pf.first_line_indent = Pt(20)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.widow_control = True

    h1 = doc.styles["Heading 1"]
    h1.font.name = FONT_SANS
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SANS)
    h1.font.size = Pt(15)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(9)
    h1.paragraph_format.space_after = Pt(5)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.first_line_indent = None

    h2 = doc.styles["Heading 2"]
    h2.font.name = FONT_SANS
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SANS)
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.space_after = Pt(2)
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.first_line_indent = None


def add_cover(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    r = p.add_run(title)
    set_run_font(r, 22, True, True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(22)
    r = p2.add_run("项目申请书（内容验证稿）")
    set_run_font(r, 16, True, True)

    tbl = doc.add_table(rows=5, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Cm(4.2)
    tbl.columns[1].width = Cm(8.8)
    entries = [
        ("项目名称", title),
        ("申报单位", "待补充"),
        ("项目负责人", "待补充"),
        ("项目周期", "待正式指南与申报要求确定"),
        ("编制日期", "2026年7月"),
    ]
    for row, (k, v) in zip(tbl.rows, entries):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(0.95)
        for c in row.cells:
            set_cell_margins(c, 80, 100, 80, 100)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[0].text = k
        row.cells[1].text = v
        for j, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.first_line_indent = None
                for run in para.runs:
                    set_run_font(run, 11, bold=(j == 0), sans=(j == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(28)
    rn = note.add_run("说明：正式申报指南、模板、团队、经费与周期信息到位后，需进行最终适配。")
    set_run_font(rn, 9.5, False)


def add_header_footer(section, title: str, page_number: bool) -> None:
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = header.add_run(title)
    set_run_font(hr, 8.5, False, color="666666")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if page_number:
        add_page_field(footer)


def add_table(doc: Document, spec: dict) -> None:
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.keep_with_next = True
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(2)
    r = caption.add_run(spec["caption"])
    set_run_font(r, 9, True)
    headers = spec["headers"]
    rows = spec["rows"]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_repeat_table_header(table.rows[0])
    for j, val in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_shading(cell, "E7E7E7")
        cell.text = val
        set_cell_margins(cell, 35, 45, 35, 45)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.line_spacing = Pt(12)
            for rr in p.runs:
                set_run_font(rr, 7.6, True, True)
    for i, vals in enumerate(rows, start=1):
        row = table.rows[i]
        cant_split_row(row)
        for j, val in enumerate(vals):
            cell = row.cells[j]
            cell.text = val
            set_cell_margins(cell, 30, 40, 30, 40)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = None
                p.paragraph_format.line_spacing = Pt(11)
                p.paragraph_format.space_after = Pt(0)
                for rr in p.runs:
                    set_run_font(rr, 7.2)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def add_figure(doc: Document, spec: dict, image_path: Path) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(16.0))
    doc_pr_nodes = run._r.xpath(".//wp:docPr")
    if doc_pr_nodes:
        doc_pr_nodes[0].set("title", spec["caption"])
        doc_pr_nodes[0].set("descr", spec["caption"] + "。" + spec["footer"])
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(2)
    rc = cap.add_run(spec["caption"])
    set_run_font(rc, 9, True)


def add_reference(doc: Document, line: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.72)
    p.paragraph_format.first_line_indent = Cm(-0.72)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(14)
    p.paragraph_format.space_after = Pt(0.5)
    r = p.add_run(line)
    set_run_font(r, 8.5)


def parse_markdown(path: Path) -> tuple[str, list[tuple[str, str]]]:
    title = ""
    blocks: list[tuple[str, str]] = []
    buffer: list[str] = []
    def flush():
        nonlocal buffer
        text = " ".join(s.strip() for s in buffer if s.strip()).strip()
        if text:
            blocks.append(("p", text))
        buffer = []
    for raw in path.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if line.startswith("# "):
            flush()
            heading = line[2:].strip()
            if not title:
                title = heading
            else:
                blocks.append(("h1", heading))
        elif line.startswith("## "):
            flush()
            blocks.append(("h2", line[3:].strip()))
        elif line.startswith("> **"):
            flush()
            marker = line.replace("> **", "", 1).rsplit("**", 1)[0].strip()
            blocks.append(("visual", marker))
        elif not line.strip():
            flush()
        else:
            buffer.append(line)
    flush()
    return title, blocks


def build_docx(md_path: Path, out_docx: Path, asset_dir: Path) -> dict:
    title, blocks = parse_markdown(md_path)
    asset_dir.mkdir(parents=True, exist_ok=True)
    for key, spec in FIGURE_SPECS.items():
        draw_flow_figure(spec, asset_dir / f"{key}.png")

    doc = Document()
    configure_styles(doc)
    add_cover(doc, title)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    body_section.page_width = Cm(21)
    body_section.page_height = Cm(29.7)
    body_section.top_margin = Cm(1.9)
    body_section.bottom_margin = Cm(1.8)
    body_section.left_margin = Cm(2.15)
    body_section.right_margin = Cm(2.05)
    body_section.header_distance = Cm(0.8)
    body_section.footer_distance = Cm(0.8)
    body_section.header.is_linked_to_previous = False
    body_section.footer.is_linked_to_previous = False
    restart_page_number(body_section, 1)
    add_header_footer(body_section, title, True)

    in_refs = False
    ref_start_marker = ""
    h1_count = 0
    for kind, value in blocks:
        if kind == "h1":
            if value == "参考文献":
                doc.add_page_break()
                in_refs = True
                p = doc.add_paragraph(style="Heading 1")
                p.paragraph_format.page_break_before = False
                r = p.add_run("参考文献")
                set_run_font(r, 15, True, True)
                continue
            h1_count += 1
            num = CHAPTER_NUMBERS.get(value, str(h1_count))
            p = doc.add_paragraph(style="Heading 1")
            r = p.add_run(f"{num}、{value}")
            set_run_font(r, 15, True, True)
        elif kind == "h2":
            p = doc.add_paragraph(style="Heading 2")
            r = p.add_run(value)
            set_run_font(r, 12, True, True)
        elif kind == "visual":
            key = value.split("：", 1)[0].strip()
            if key in FIGURE_SPECS:
                add_figure(doc, FIGURE_SPECS[key], asset_dir / f"{key}.png")
            elif key in TABLE_SPECS:
                add_table(doc, TABLE_SPECS[key])
        elif kind == "p":
            if in_refs:
                references = [item.strip() for item in re.split(r"(?=\[\d+\])", value) if item.strip()]
                for reference in references:
                    add_reference(doc, reference)
            else:
                p = doc.add_paragraph()
                r = p.add_run(value)
                set_run_font(r, 10.5)

    # Prevent title-page footer inherited blank line from showing a page number.
    doc.sections[0].header.is_linked_to_previous = False
    doc.sections[0].footer.is_linked_to_previous = False
    doc.sections[0].header.paragraphs[0].clear()
    doc.sections[0].footer.paragraphs[0].clear()

    # Core properties.
    doc.core_properties.title = title
    doc.core_properties.subject = "项目申请书"
    doc.core_properties.author = "Proposal Agent System"
    doc.core_properties.keywords = "人机协同决策, 项目申请书"
    doc.core_properties.comments = "由智能体工作流冻结正文后进行确定性排版导出。"
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx)
    return {
        "title": title,
        "chapter_count": h1_count,
        "table_count": len(TABLE_SPECS),
        "figure_count": len(FIGURE_SPECS),
        "docx_sha256": sha256(out_docx),
        "docx_size": out_docx.stat().st_size,
    }


def convert_pdf(docx_path: Path, out_pdf: Path) -> dict:
    out_pdf.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="stage8-lo-") as profile:
        cmd = [
            shutil.which("libreoffice") or "libreoffice",
            f"-env:UserInstallation={Path(profile).resolve().as_uri()}",
            "--headless", "--convert-to", "pdf:writer_pdf_Export",
            "--outdir", str(out_pdf.parent), str(docx_path),
        ]
        cp = subprocess.run(cmd, capture_output=True, text=True, timeout=240)
    generated = out_pdf.parent / f"{docx_path.stem}.pdf"
    if cp.returncode != 0 or not generated.exists():
        raise RuntimeError(f"LibreOffice conversion failed: {cp.stderr or cp.stdout}")
    if generated.resolve() != out_pdf.resolve():
        generated.replace(out_pdf)
    reader = PdfReader(str(out_pdf))
    return {
        "pdf_sha256": sha256(out_pdf),
        "pdf_size": out_pdf.stat().st_size,
        "page_count": len(reader.pages),
        "stdout": cp.stdout,
        "stderr": cp.stderr,
    }


def page_locations(pdf_path: Path) -> dict:
    reader = PdfReader(str(pdf_path))
    first_body = None
    refs = None
    chapters = {}
    for idx, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if first_body is None and "摘要与项目概览" in text:
            first_body = idx
        if refs is None and re.search(r"参考文献", text) and "[1]" in text:
            refs = idx
        for title, num in CHAPTER_NUMBERS.items():
            if f"{num}、{title}" in text and title not in chapters:
                chapters[title] = idx
    if first_body is None:
        raise RuntimeError("Cannot locate first body page")
    if refs is None:
        raise RuntimeError("Cannot locate references start page")
    body_pages = refs - first_body
    return {
        "total_pages": len(reader.pages),
        "cover_pages": first_body - 1,
        "body_start_page": first_body,
        "references_start_page": refs,
        "body_page_count": body_pages,
        "reference_page_count": len(reader.pages) - refs + 1,
        "chapter_start_pages": chapters,
    }


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", required=True)
    ap.add_argument("--out-dir", required=True)
    args = ap.parse_args()
    md_path = Path(args.input).resolve()
    out_dir = Path(args.out_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    docx_path = out_dir / "人机协同决策优势冲刺关键技术研究_申请书.docx"
    pdf_path = out_dir / "人机协同决策优势冲刺关键技术研究_申请书.pdf"
    metadata = build_docx(md_path, docx_path, out_dir / "assets")
    metadata.update(convert_pdf(docx_path, pdf_path))
    metadata.update(page_locations(pdf_path))
    metadata["input_sha256"] = sha256(md_path)
    metadata["input_file"] = str(md_path)
    (out_dir / "stage8_export_metadata.json").write_text(
        json.dumps(metadata, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(json.dumps(metadata, ensure_ascii=False, indent=2))

if __name__ == "__main__":
    main()
